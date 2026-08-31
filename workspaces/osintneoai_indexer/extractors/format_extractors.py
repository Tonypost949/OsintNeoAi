"""
OsintNeoAi Indexer: Format-Specific Evidentiary Document Extractors
Path: C:\\OsintNeoAi\\workspaces\\osintneoai_indexer\\extractors\\format_extractors.py
Milestone: M2 (Deep Text Extraction & OCR Engine) — Features 5, 6, 7

Provides specialized, memory-bounded format extractors:
1. TiffExtractor: Multi-frame TIFF streaming with Pillow & RapidOCR
2. HtmlDocumentParser: Structured HTML -> Markdown with tables & metadata
3. DocxExtractor: OOXML python-docx parser with tables, headers/footers & comments.xml
4. ImageExtractor: EXIF transposition, two-pass OpenCV enhancement & OCR
5. TextExtractor: Multi-encoding text reader, CSV sniffer & JSON formatter
"""

from __future__ import annotations

import csv
import gc
import io
import json
import logging
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, BinaryIO, Dict, List, Optional, Sequence, Tuple, Union

import chardet
import cv2
import docx
from docx.table import Table
import numpy as np
from PIL import Image, ImageOps, ImageSequence

import lxml.etree
import lxml.html

from config import OCR_CONFIDENCE_THRESHOLD, IndexerConfig
from extractors.image_enhancer import EnhancementProfile, ImageEnhancer
from extractors.ocr_engine import OCREngine, RapidOCREngine

logger = logging.getLogger("osintneoai.extractors.format_extractors")


# ==============================================================================
# 1. Multi-Page / Multi-Frame TIFF Extractor
# ==============================================================================

@dataclass
class TiffPageResult:
    """Individual page extraction telemetry for TIFF documents."""
    page_number: int
    text: str
    confidence: float
    dimensions: Tuple[int, int]
    original_mode: str
    line_count: int


@dataclass
class TiffExtractionResult:
    """Aggregated multi-page TIFF extraction result."""
    full_text: str
    page_count: int
    average_confidence: float
    pages: List[TiffPageResult]
    metadata: Dict[str, Any]
    ocr_engine_used: str = "rapidocr_onnx_tiff"


class TiffExtractor:
    """
    High-performance, memory-bounded multi-page TIFF extractor.
    Streams frames from binary streams or file paths, normalizes color spaces,
    and executes RapidOCR ONNX inference per frame with explicit deallocations.
    """

    def __init__(
        self,
        ocr_engine: Optional[OCREngine] = None,
        confidence_threshold: float = OCR_CONFIDENCE_THRESHOLD,
    ) -> None:
        self.ocr_engine = ocr_engine or OCREngine.get_instance()
        self.confidence_threshold = confidence_threshold

    def extract_from_stream(
        self,
        stream: BinaryIO,
        source_uri: str = "stream://tiff",
    ) -> TiffExtractionResult:
        """Extracts structured text and OCR transcripts from a binary stream."""
        try:
            if not stream.seekable():
                stream_bytes = stream.read()
                img = Image.open(io.BytesIO(stream_bytes))
            else:
                stream.seek(0)
                img = Image.open(stream)
        except Exception as e:
            logger.error(f"Failed to open TIFF stream from {source_uri}: {e}")
            return TiffExtractionResult(
                full_text=f"[ERROR: Corrupt or unreadable TIFF archive: {e}]",
                page_count=0,
                average_confidence=0.0,
                pages=[],
                metadata={"error": str(e), "source_uri": source_uri},
            )

        return self._process_image(img, source_uri)

    def extract_from_path(self, file_path: Union[str, Path]) -> TiffExtractionResult:
        """Extracts structured text from a local TIFF file."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"TIFF file not found: {path}")

        try:
            with Image.open(path) as img:
                return self._process_image(img, str(path))
        except Exception as e:
            logger.error(f"Error processing TIFF file {path}: {e}")
            return TiffExtractionResult(
                full_text=f"[ERROR: Unreadable TIFF file {path.name}: {e}]",
                page_count=0,
                average_confidence=0.0,
                pages=[],
                metadata={"error": str(e), "file_path": str(path)},
            )

    def _process_image(self, img: Image.Image, source_uri: str) -> TiffExtractionResult:
        """Internal frame-by-frame OCR processing loop."""
        total_frames = getattr(img, "n_frames", 1)
        pages_results: List[TiffPageResult] = []
        page_texts: List[str] = []
        total_conf_sum = 0.0
        valid_conf_pages = 0

        for frame_idx, frame in enumerate(ImageSequence.Iterator(img)):
            page_num = frame_idx + 1
            orig_mode = frame.mode
            orig_size = frame.size

            try:
                # Convert all modes (1, L, P, RGBA, CMYK, I;16) to standard 8-bit RGB
                frame_rgb = frame.convert("RGB")
                np_arr = np.array(frame_rgb, dtype=np.uint8)

                # Execute OCR
                ocr_res = self.ocr_engine.ocr_image(np_arr, page_number=page_num)
                page_text = ocr_res.full_text.strip()
                page_conf = ocr_res.avg_confidence

                if page_conf > 0.0:
                    total_conf_sum += page_conf
                    valid_conf_pages += 1

                pages_results.append(
                    TiffPageResult(
                        page_number=page_num,
                        text=page_text,
                        confidence=page_conf,
                        dimensions=orig_size,
                        original_mode=orig_mode,
                        line_count=len(ocr_res.lines),
                    )
                )

                if page_text:
                    page_texts.append(f"--- [TIFF Page {page_num}/{total_frames}] ---\n{page_text}")
                else:
                    page_texts.append(f"--- [TIFF Page {page_num}/{total_frames}] ---\n[Blank or unreadable page]")

            except Exception as frame_err:
                logger.warning(f"Error processing frame {page_num} of {source_uri}: {frame_err}")
                pages_results.append(
                    TiffPageResult(
                        page_number=page_num,
                        text="",
                        confidence=0.0,
                        dimensions=orig_size,
                        original_mode=orig_mode,
                        line_count=0,
                    )
                )
            finally:
                # Explicit memory reclamation for high-res frames
                if "frame_rgb" in locals():
                    del frame_rgb
                if "np_arr" in locals():
                    del np_arr
                if page_num % 5 == 0:
                    gc.collect()

        avg_confidence = (total_conf_sum / valid_conf_pages) if valid_conf_pages > 0 else 0.0
        full_text_body = "\n\n\f\n\n".join(page_texts)

        return TiffExtractionResult(
            full_text=full_text_body,
            page_count=total_frames,
            average_confidence=round(avg_confidence, 4),
            pages=pages_results,
            metadata={
                "source_uri": source_uri,
                "total_pages": total_frames,
                "average_confidence": round(avg_confidence, 4),
            },
        )


# ==============================================================================
# 2. Structured HTML Document Parser
# ==============================================================================

@dataclass
class HtmlExtractionResult:
    """Structured extraction output for HTML documents."""
    text: str
    title: Optional[str]
    meta_tags: Dict[str, str]
    links: List[Dict[str, str]]
    email_addresses: List[str]
    metadata: Dict[str, Any]
    ocr_engine_used: str = "lxml_html_parser"


class StdlibHtmlFallbackParser(HTMLParser):
    """Robust standard library fallback parser for malformed HTML."""

    def __init__(self) -> None:
        super().__init__()
        self.text_parts: List[str] = []
        self.title: Optional[str] = None
        self._in_title: bool = False
        self._skip_tags = {"script", "style", "noscript", "svg", "iframe", "canvas"}
        self._current_skip: int = 0

    def handle_starttag(self, tag: str, attrs: List[Tuple[str, Optional[str]]]) -> None:
        t = tag.lower()
        if t in self._skip_tags:
            self._current_skip += 1
        elif t == "title":
            self._in_title = True
        elif t in ("h1", "h2", "h3", "h4", "h5", "h6"):
            try:
                lvl = int(t[1])
            except Exception:
                lvl = 1
            self.text_parts.append(f"\n\n{'#' * lvl} ")
        elif t == "p":
            self.text_parts.append("\n\n")
        elif t == "br":
            self.text_parts.append("\n")
        elif t == "li":
            self.text_parts.append("\n- ")

    def handle_endtag(self, tag: str) -> None:
        t = tag.lower()
        if t in self._skip_tags and self._current_skip > 0:
            self._current_skip -= 1
        elif t == "title":
            self._in_title = False
        elif t in ("p", "h1", "h2", "h3", "h4", "h5", "h6"):
            self.text_parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._current_skip > 0:
            return
        if self._in_title:
            self.title = (self.title or "") + data
        cleaned = " ".join(data.split())
        if cleaned:
            self.text_parts.append(cleaned + " ")

    def get_text(self) -> str:
        return re.sub(r"\n{3,}", "\n\n", "".join(self.text_parts)).strip()


class HtmlDocumentParser:
    """
    High-performance HTML document parser converting evidentiary web records
    into clean, structured Markdown text while preserving tables and metadata.
    """

    STRIP_TAGS = (
        "script", "style", "noscript", "iframe", "svg",
        "canvas", "template", "nav", "footer", "header"
    )

    def extract_from_stream(
        self,
        stream: BinaryIO,
        source_uri: str = "stream://html",
    ) -> HtmlExtractionResult:
        """Parses HTML document from raw binary stream."""
        raw_bytes = stream.read()
        return self.extract_from_bytes(raw_bytes, source_uri)

    def extract_from_path(self, file_path: Union[str, Path]) -> HtmlExtractionResult:
        """Parses HTML document from local file path."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"HTML file not found: {path}")
        raw_bytes = path.read_bytes()
        return self.extract_from_bytes(raw_bytes, str(path))

    def extract_from_bytes(
        self,
        raw_bytes: bytes,
        source_uri: str = "memory://html",
    ) -> HtmlExtractionResult:
        """Core decoding and parsing routine."""
        if not raw_bytes.strip():
            return HtmlExtractionResult(
                text="",
                title=None,
                meta_tags={},
                links=[],
                email_addresses=[],
                metadata={"source_uri": source_uri, "empty_payload": True},
            )

        html_text = self._decode_bytes(raw_bytes)
        try:
            return self._parse_lxml(html_text, source_uri)
        except Exception as lxml_err:
            logger.warning(f"lxml parsing failed for {source_uri}: {lxml_err}. Falling back to stdlib HTMLParser.")
            return self._parse_stdlib(html_text, source_uri)

    def _decode_bytes(self, raw_bytes: bytes) -> str:
        """Multi-tiered encoding resolution."""
        # 1. UTF-8 with BOM or plain UTF-8
        try:
            return raw_bytes.decode("utf-8-sig")
        except UnicodeDecodeError:
            pass

        # 2. Check meta charset tag in first 2048 bytes
        header_sample = raw_bytes[:2048].decode("ascii", errors="ignore")
        match = re.search(r'charset=["\']?([a-zA-Z0-9_-]+)', header_sample, re.IGNORECASE)
        if match:
            charset = match.group(1).lower()
            try:
                return raw_bytes.decode(charset)
            except (UnicodeDecodeError, LookupError):
                pass

        # 3. Chardet heuristic
        detected = chardet.detect(raw_bytes[:65536])
        if detected and detected.get("encoding"):
            try:
                return raw_bytes.decode(detected["encoding"])
            except (UnicodeDecodeError, LookupError):
                pass

        # 4. Fallback to windows-1252 / latin-1 with replacement
        return raw_bytes.decode("latin-1", errors="replace")

    def _parse_lxml(self, html_text: str, source_uri: str) -> HtmlExtractionResult:
        """Structured DOM parsing using lxml."""
        root = lxml.html.fromstring(html_text)

        # 1. Sanitize elements
        lxml.etree.strip_elements(root, *self.STRIP_TAGS)

        # 2. Extract Document Title
        title = root.findtext(".//title")
        if title:
            title = " ".join(title.split()).strip()

        # 3. Extract Meta Tags
        meta_tags: Dict[str, str] = {}
        for meta in root.xpath(".//meta"):
            name = meta.get("name") or meta.get("property") or meta.get("http-equiv")
            content = meta.get("content")
            if name and content:
                meta_tags[name.lower().strip()] = " ".join(content.split()).strip()

        # 4. Extract Links & Mailto Addresses
        links: List[Dict[str, str]] = []
        emails: List[str] = []
        for a_tag in root.xpath(".//a"):
            href = a_tag.get("href")
            link_text = " ".join(a_tag.text_content().split()).strip()
            if href:
                href_clean = href.strip()
                if href_clean.lower().startswith("mailto:"):
                    email_addr = href_clean[7:].split("?")[0].strip()
                    if email_addr and email_addr not in emails:
                        emails.append(email_addr)
                else:
                    links.append({"text": link_text, "url": href_clean})

        # 5. Extract Structured Markdown Text Body
        blocks: List[str] = []
        if title:
            blocks.append(f"# {title}")

        body_elem = root.body if root.body is not None else root
        for elem in body_elem.iterchildren():
            tag = elem.tag.lower() if isinstance(elem.tag, str) else ""

            if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                try:
                    level = int(tag[1])
                except Exception:
                    level = 1
                heading_text = " ".join(elem.text_content().split()).strip()
                if heading_text:
                    blocks.append(f"{'#' * level} {heading_text}")

            elif tag == "p":
                p_text = " ".join(elem.text_content().split()).strip()
                if p_text:
                    blocks.append(p_text)

            elif tag == "table":
                md_table = self._format_table(elem)
                if md_table:
                    blocks.append(md_table)

            elif tag in ("ul", "ol"):
                list_items = []
                for li in elem.xpath(".//li"):
                    li_text = " ".join(li.text_content().split()).strip()
                    if li_text:
                        list_items.append(f"- {li_text}")
                if list_items:
                    blocks.append("\n".join(list_items))

            elif tag in ("blockquote", "section", "article", "div"):
                div_text = " ".join(elem.text_content().split()).strip()
                if div_text:
                    blocks.append(div_text)

        if not blocks:
            raw_content = " ".join(root.text_content().split()).strip()
            blocks.append(raw_content)

        markdown_body = "\n\n".join(blocks).strip()

        return HtmlExtractionResult(
            text=markdown_body,
            title=title,
            meta_tags=meta_tags,
            links=links[:50],
            email_addresses=emails,
            metadata={
                "source_uri": source_uri,
                "title": title,
                "meta_tags": meta_tags,
                "link_count": len(links),
                "email_count": len(emails),
            },
            ocr_engine_used="lxml_html_parser",
        )

    def _format_table(self, table_elem: lxml.html.HtmlElement) -> Optional[str]:
        """Converts an HTML table element to Markdown table text."""
        rows_data: List[List[str]] = []
        for tr in table_elem.xpath(".//tr"):
            cells = [" ".join(c.text_content().split()).replace("|", "\\|").strip() for c in tr.xpath(".//th | .//td")]
            if any(cells):
                rows_data.append(cells)

        if not rows_data:
            return None

        col_count = max(len(r) for r in rows_data)
        if col_count == 0:
            return None

        padded_rows = [r + [""] * (col_count - len(r)) for r in rows_data]
        header = "| " + " | ".join(padded_rows[0]) + " |"
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        body = ["| " + " | ".join(r) + " |" for r in padded_rows[1:]]

        return "\n".join([header, separator] + body)

    def _parse_stdlib(self, html_text: str, source_uri: str) -> HtmlExtractionResult:
        """Fallback parsing when lxml fails."""
        parser = StdlibHtmlFallbackParser()
        parser.feed(html_text)
        text = parser.get_text()
        return HtmlExtractionResult(
            text=text,
            title=parser.title,
            meta_tags={},
            links=[],
            email_addresses=[],
            metadata={"source_uri": source_uri, "fallback": True},
            ocr_engine_used="stdlib_html_parser",
        )


# ==============================================================================
# 3. DOCX Document Extractor
# ==============================================================================

@dataclass
class DocxComment:
    """Forensic annotation/comment extracted from DOCX internal XML."""
    author: str
    date: Optional[str]
    text: str


@dataclass
class DocxExtractionResult:
    """Structured extraction output for Microsoft Word documents."""
    text: str
    title: Optional[str]
    author: Optional[str]
    created_date: Optional[str]
    modified_date: Optional[str]
    comments: List[DocxComment]
    metadata: Dict[str, Any]
    ocr_engine_used: str = "docx_native_parser"


class DocxExtractor:
    """
    Forensic DOCX extractor extracting headings, paragraphs, tables,
    headers, footers, comments, and core metadata.
    """

    def extract_from_stream(
        self,
        stream: BinaryIO,
        source_uri: str = "stream://docx",
    ) -> DocxExtractionResult:
        """Parses DOCX document from raw binary stream."""
        raw_bytes = stream.read()
        return self.extract_from_bytes(raw_bytes, source_uri)

    def extract_from_path(self, file_path: Union[str, Path]) -> DocxExtractionResult:
        """Parses DOCX document from local file path."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")
        raw_bytes = path.read_bytes()
        return self.extract_from_bytes(raw_bytes, str(path))

    def extract_from_bytes(
        self,
        raw_bytes: bytes,
        source_uri: str = "memory://docx",
    ) -> DocxExtractionResult:
        """Core parsing implementation using python-docx and zip inspection."""
        if not raw_bytes.strip():
            return DocxExtractionResult(
                text="",
                title=None,
                author=None,
                created_date=None,
                modified_date=None,
                comments=[],
                metadata={"source_uri": source_uri, "empty_payload": True},
            )

        bio = io.BytesIO(raw_bytes)
        try:
            doc = docx.Document(bio)
        except Exception as doc_err:
            logger.error(f"Failed to parse DOCX document from {source_uri}: {doc_err}")
            return DocxExtractionResult(
                text=f"[ERROR: Corrupt or encrypted DOCX document: {doc_err}]",
                title=None,
                author=None,
                created_date=None,
                modified_date=None,
                comments=[],
                metadata={"error": str(doc_err), "source_uri": source_uri},
            )

        # 1. Extract Core Properties
        props = doc.core_properties
        title = props.title.strip() if props.title else None
        author = props.author.strip() if props.author else None
        created = props.created.isoformat() if props.created else None
        modified = props.modified.isoformat() if props.modified else None

        # 2. Extract Forensic Comments from zip container
        comments = self._extract_comments(raw_bytes)

        # 3. Extract Section Headers & Footers
        header_footer_texts = self._extract_headers_footers(doc)

        # 4. Extract Paragraphs & Headings
        blocks: List[str] = []
        if title:
            blocks.append(f"# {title}")

        if header_footer_texts:
            blocks.append(f"<!-- Headers/Footers: {', '.join(header_footer_texts)} -->")

        for p in doc.paragraphs:
            p_text = p.text.strip()
            if not p_text:
                continue

            style_name = p.style.name if p.style else ""
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                except Exception:
                    level = 1
                blocks.append(f"{'#' * level} {p_text}")
            elif "List" in style_name:
                blocks.append(f"- {p_text}")
            else:
                blocks.append(p_text)

        # 5. Extract Tables
        for table in doc.tables:
            md_table = self._format_table(table)
            if md_table:
                blocks.append(md_table)

        # 6. Append Comments Section if present
        if comments:
            comment_blocks = ["\n### Forensic Document Annotations / Comments"]
            for c in comments:
                date_str = f" ({c.date})" if c.date else ""
                comment_blocks.append(f"- **{c.author}**{date_str}: {c.text}")
            blocks.append("\n".join(comment_blocks))

        full_text = "\n\n".join(blocks).strip()

        return DocxExtractionResult(
            text=full_text,
            title=title,
            author=author,
            created_date=created,
            modified_date=modified,
            comments=comments,
            metadata={
                "source_uri": source_uri,
                "title": title,
                "author": author,
                "created": created,
                "modified": modified,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "comment_count": len(comments),
            },
        )

    def _extract_headers_footers(self, doc: docx.Document) -> List[str]:
        """Extracts unique text strings from document section headers and footers."""
        hf_texts: List[str] = []
        seen = set()

        for section in doc.sections:
            if hasattr(section, "header") and section.header:
                for p in section.header.paragraphs:
                    t = p.text.strip()
                    if t and t not in seen:
                        seen.add(t)
                        hf_texts.append(t)
            if hasattr(section, "footer") and section.footer:
                for p in section.footer.paragraphs:
                    t = p.text.strip()
                    if t and t not in seen:
                        seen.add(t)
                        hf_texts.append(t)

        return hf_texts

    def _format_table(self, table: Table) -> Optional[str]:
        """Formats a docx Table into a Markdown grid table."""
        rows_data: List[List[str]] = []
        for row in table.rows:
            cells = [" ".join(cell.text.strip().split()).replace("|", "\\|") for cell in row.cells]
            if any(cells):
                rows_data.append(cells)

        if not rows_data:
            return None

        col_count = max(len(r) for r in rows_data)
        if col_count == 0:
            return None

        padded_rows = [r + [""] * (col_count - len(r)) for r in rows_data]
        header = "| " + " | ".join(padded_rows[0]) + " |"
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        body = ["| " + " | ".join(r) + " |" for r in padded_rows[1:]]

        return "\n".join([header, separator] + body)

    def _extract_comments(self, raw_bytes: bytes) -> List[DocxComment]:
        """Parses word/comments.xml inside the OOXML zip package if present."""
        comments: List[DocxComment] = []
        try:
            with zipfile.ZipFile(io.BytesIO(raw_bytes), "r") as zf:
                if "word/comments.xml" in zf.namelist():
                    comments_xml = zf.read("word/comments.xml")
                    root = ET.fromstring(comments_xml)
                    # Namespace for WordprocessingML
                    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    for comment_elem in root.findall(".//w:comment", ns):
                        author = comment_elem.get(f"{{{ns['w']}}}author", "Unknown Author")
                        date_str = comment_elem.get(f"{{{ns['w']}}}date")
                        text_nodes = comment_elem.findall(".//w:t", ns)
                        comment_text = "".join(t.text for t in text_nodes if t.text).strip()
                        if comment_text:
                            comments.append(DocxComment(author=author, date=date_str, text=comment_text))
        except Exception as e:
            logger.debug(f"Could not extract comments from docx: {e}")
        return comments


# ==============================================================================
# 4. Direct Raster Image Extractor
# ==============================================================================

@dataclass
class ImageExtractionResult:
    """Structured extraction output for single/multi-frame raster images."""
    text: str
    confidence: float
    dimensions: Tuple[int, int]
    format: str
    exif_data: Dict[str, Any]
    metadata: Dict[str, Any]
    ocr_engine_used: str = "rapidocr_onnx"


class ImageExtractor:
    """
    Direct image extraction engine for PNG, JPG, JPEG, WEBP, BMP, GIF.
    Performs EXIF rotation correction, two-pass OCR with OpenCV enhancement fallback.
    """

    def __init__(
        self,
        ocr_engine: Optional[OCREngine] = None,
        image_enhancer: Optional[ImageEnhancer] = None,
        confidence_threshold: float = OCR_CONFIDENCE_THRESHOLD,
    ) -> None:
        self.ocr_engine = ocr_engine or OCREngine.get_instance()
        self.image_enhancer = image_enhancer or ImageEnhancer()
        self.confidence_threshold = confidence_threshold

    def extract_from_stream(
        self,
        stream: BinaryIO,
        source_uri: str = "stream://image",
    ) -> ImageExtractionResult:
        """Extracts text from image binary stream."""
        raw_bytes = stream.read()
        return self.extract_from_bytes(raw_bytes, source_uri)

    def extract_from_path(self, file_path: Union[str, Path]) -> ImageExtractionResult:
        """Extracts text from local image file."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Image file not found: {path}")
        raw_bytes = path.read_bytes()
        return self.extract_from_bytes(raw_bytes, str(path))

    def extract_from_bytes(
        self,
        raw_bytes: bytes,
        source_uri: str = "memory://image",
    ) -> ImageExtractionResult:
        """Core decoding, EXIF orientation correction, enhancement, and OCR."""
        if not raw_bytes.strip():
            return ImageExtractionResult(
                text="",
                confidence=0.0,
                dimensions=(0, 0),
                format="UNKNOWN",
                exif_data={},
                metadata={"source_uri": source_uri, "empty_payload": True},
            )

        # 1. Load with PIL to extract EXIF and apply transposition
        exif_dict: Dict[str, Any] = {}
        img_format = "UNKNOWN"
        try:
            with Image.open(io.BytesIO(raw_bytes)) as pil_img:
                img_format = pil_img.format or "UNKNOWN"
                # Extract EXIF tags
                if hasattr(pil_img, "getexif"):
                    exif = pil_img.getexif()
                    if exif:
                        for tag_id, value in exif.items():
                            exif_dict[str(tag_id)] = str(value)
                # Correct orientation
                transposed = ImageOps.exif_transpose(pil_img)
                rgb_img = transposed.convert("RGB")
                np_arr = np.array(rgb_img, dtype=np.uint8)
                dimensions = (pil_img.width, pil_img.height)
        except Exception as pil_err:
            logger.warning(f"PIL decode failed for {source_uri}: {pil_err}. Falling back to cv2.imdecode.")
            nparr = np.frombuffer(raw_bytes, np.uint8)
            np_arr = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            if np_arr is None:
                return ImageExtractionResult(
                    text=f"[ERROR: Unreadable image format: {pil_err}]",
                    confidence=0.0,
                    dimensions=(0, 0),
                    format="UNKNOWN",
                    exif_data={},
                    metadata={"error": str(pil_err), "source_uri": source_uri},
                )
            dimensions = (np_arr.shape[1], np_arr.shape[0])

        # 2. First-Pass OCR
        ocr_res = self.ocr_engine.ocr_image(np_arr, page_number=1)
        method_used = "rapidocr_onnx"

        # 3. Second-Pass Enhancement Fallback if low confidence or no lines
        if not ocr_res.lines or ocr_res.avg_confidence < self.confidence_threshold:
            enhanced_np = self.image_enhancer.enhance(np_arr, profile=EnhancementProfile.HEAVY)
            enhanced_res = self.ocr_engine.ocr_image(enhanced_np, page_number=1)

            if len(enhanced_res.full_text) > len(ocr_res.full_text) or enhanced_res.avg_confidence > ocr_res.avg_confidence:
                ocr_res = enhanced_res
                method_used = "rapidocr_enhanced"
            del enhanced_np

        del np_arr

        return ImageExtractionResult(
            text=ocr_res.full_text,
            confidence=ocr_res.avg_confidence,
            dimensions=dimensions,
            format=img_format,
            exif_data=exif_dict,
            metadata={
                "source_uri": source_uri,
                "detection_time_sec": ocr_res.detection_time_sec,
                "recognition_time_sec": ocr_res.recognition_time_sec,
                "total_time_sec": ocr_res.total_time_sec,
            },
            ocr_engine_used=method_used,
        )


# ==============================================================================
# 5. Plaintext & Structured Data Extractor
# ==============================================================================

@dataclass
class TextExtractionResult:
    """Structured extraction output for plaintext, CSV, TSV, JSON, and XML."""
    text: str
    encoding_used: str
    format_detected: str
    metadata: Dict[str, Any]
    ocr_engine_used: str = "plaintext_reader"


class TextExtractor:
    """
    Multi-encoding resilient text extractor handling .txt, .md, .csv, .tsv,
    .json, .jsonl, .xml, and .yaml.
    """

    ENCODING_LADDER = ["utf-8-sig", "utf-8", "utf-16", "windows-1252", "latin-1"]

    def extract_from_stream(
        self,
        stream: BinaryIO,
        source_uri: str = "stream://text",
        mime_type: Optional[str] = None,
    ) -> TextExtractionResult:
        """Extracts text from binary stream."""
        raw_bytes = stream.read()
        return self.extract_from_bytes(raw_bytes, source_uri, mime_type)

    def extract_from_path(
        self,
        file_path: Union[str, Path],
        mime_type: Optional[str] = None,
    ) -> TextExtractionResult:
        """Extracts text from local file."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Text file not found: {path}")
        raw_bytes = path.read_bytes()
        return self.extract_from_bytes(raw_bytes, str(path), mime_type)

    def extract_from_bytes(
        self,
        raw_bytes: bytes,
        source_uri: str = "memory://text",
        mime_type: Optional[str] = None,
    ) -> TextExtractionResult:
        """Decodes raw bytes and formats according to detected or specified format."""
        if not raw_bytes.strip():
            return TextExtractionResult(
                text="",
                encoding_used="utf-8",
                format_detected="empty",
                metadata={"source_uri": source_uri, "empty_payload": True},
            )

        decoded_text, encoding = self._decode(raw_bytes)

        # Detect format
        ext = os.path.splitext(source_uri.lower())[1]
        if ext in (".csv", ".tsv") or mime_type in ("text/csv", "text/tab-separated-values"):
            return self._parse_csv(decoded_text, encoding, source_uri)
        elif ext in (".json", ".jsonl", ".ndjson") or mime_type in ("application/json", "application/x-ndjson"):
            return self._parse_json(decoded_text, encoding, source_uri)
        else:
            return TextExtractionResult(
                text=decoded_text.strip(),
                encoding_used=encoding,
                format_detected="plaintext",
                metadata={"source_uri": source_uri, "character_count": len(decoded_text)},
                ocr_engine_used="plaintext_reader",
            )

    def _decode(self, raw_bytes: bytes) -> Tuple[str, str]:
        """Tries encoding ladder before falling back to chardet and replacement."""
        for enc in self.ENCODING_LADDER:
            try:
                return raw_bytes.decode(enc), enc
            except (UnicodeDecodeError, LookupError):
                continue

        detected = chardet.detect(raw_bytes[:65536])
        if detected and detected.get("encoding"):
            try:
                return raw_bytes.decode(detected["encoding"]), detected["encoding"]
            except (UnicodeDecodeError, LookupError):
                pass

        return raw_bytes.decode("latin-1", errors="replace"), "latin-1-replace"

    def _parse_csv(self, text: str, encoding: str, source_uri: str) -> TextExtractionResult:
        """Parses CSV/TSV and formats into clean Markdown table."""
        try:
            sample = text[:4096]
            sniffer = csv.Sniffer()
            delimiter = ","
            try:
                dialect = sniffer.sniff(sample)
                delimiter = dialect.delimiter
            except Exception:
                if "\t" in sample and "," not in sample:
                    delimiter = "\t"

            reader = csv.reader(io.StringIO(text), delimiter=delimiter)
            rows = [row for row in reader if any(cell.strip() for cell in row)]

            if not rows:
                return TextExtractionResult(
                    text=text.strip(),
                    encoding_used=encoding,
                    format_detected="csv",
                    metadata={"source_uri": source_uri},
                    ocr_engine_used="csv_parser",
                )

            col_count = max(len(r) for r in rows)
            padded = [r + [""] * (col_count - len(r)) for r in rows]
            header = "| " + " | ".join(c.replace("|", "\\|").strip() for c in padded[0]) + " |"
            separator = "| " + " | ".join(["---"] * col_count) + " |"
            body = ["| " + " | ".join(c.replace("|", "\\|").strip() for c in r) + " |" for r in padded[1:]]

            md_table = "\n".join([header, separator] + body)
            return TextExtractionResult(
                text=md_table,
                encoding_used=encoding,
                format_detected="csv",
                metadata={"source_uri": source_uri, "row_count": len(rows), "column_count": col_count},
                ocr_engine_used="csv_parser",
            )
        except Exception as e:
            logger.debug(f"CSV formatting error for {source_uri}: {e}")
            return TextExtractionResult(
                text=text.strip(),
                encoding_used=encoding,
                format_detected="csv_raw",
                metadata={"source_uri": source_uri, "parse_error": str(e)},
                ocr_engine_used="plaintext_reader",
            )

    def _parse_json(self, text: str, encoding: str, source_uri: str) -> TextExtractionResult:
        """Pretty-prints JSON data or NDJSON streams."""
        try:
            # Try single JSON object/array
            data = json.loads(text)
            formatted = json.dumps(data, indent=2, ensure_ascii=False)
            return TextExtractionResult(
                text=formatted,
                encoding_used=encoding,
                format_detected="json",
                metadata={"source_uri": source_uri, "type": type(data).__name__},
                ocr_engine_used="json_parser",
            )
        except json.JSONDecodeError:
            # Try newline-delimited JSON
            lines = text.strip().split("\n")
            json_objects = []
            for line in lines:
                l_str = line.strip()
                if l_str:
                    try:
                        json_objects.append(json.loads(l_str))
                    except json.JSONDecodeError:
                        pass
            if json_objects:
                formatted = "\n\n".join(json.dumps(obj, indent=2, ensure_ascii=False) for obj in json_objects)
                return TextExtractionResult(
                    text=formatted,
                    encoding_used=encoding,
                    format_detected="ndjson",
                    metadata={"source_uri": source_uri, "object_count": len(json_objects)},
                    ocr_engine_used="json_parser",
                )

        return TextExtractionResult(
            text=text.strip(),
            encoding_used=encoding,
            format_detected="json_raw",
            metadata={"source_uri": source_uri},
            ocr_engine_used="plaintext_reader",
        )
