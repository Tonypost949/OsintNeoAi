"""
OsintNeoAi Indexer — Global Pytest Fixtures & Test Harness
Path: C:\\OsintNeoAi\\workspaces\\osintneoai_indexer\\tests\\conftest.py

Provides self-contained, isolated test environments, synthetic evidentiary artifacts,
temporary SQLite databases with 100% 3NF schema DDL, mock Google Drive streams,
RFC 8785 JSON canonicalization helpers, and Merkle tree calculators.
"""

from __future__ import annotations

import email
import hashlib
import io
import json
import os
import sqlite3
import sys
import tarfile
import tempfile
import zipfile
from email.message import EmailMessage
from pathlib import Path
from typing import Any, BinaryIO, Callable, Dict, Generator, List, Optional, Tuple, Union

import docx
import numpy as np
from PIL import Image, ImageDraw, ImageFont
import pymupdf
import pytest

# Ensure both workspace root and project root are in sys.path
WORKSPACE_ROOT = Path(__file__).resolve().parent.parent
PROJECT_ROOT = WORKSPACE_ROOT.parent.parent

for p in [str(WORKSPACE_ROOT), str(PROJECT_ROOT)]:
    if p not in sys.path:
        sys.path.insert(0, p)

from config import (
    CHUNK_SIZE,
    FileCategory,
    IndexerConfig,
    get_file_category,
    get_mime_type,
    is_ignored_file,
    is_supported_file,
)
from storage.hasher import (
    StreamHasher,
    HashingReader,
    compute_bytes_sha256,
    compute_file_sha256,
    compute_stream_sha256,
    verify_file_sha256,
    verify_stream_sha256,
)
from connectors.local_crawler import IngestedArtifact, LocalCrawler, detect_mime_type
from connectors.gdrive_streamer import GDriveStreamer, GDriveResourceInfo
from connectors.mailbox_reader import MailboxReader, EmailMetadata
from extractors.document_extractor import DocumentExtractor, ExtractedRecord
from extractors.format_extractors import (
    DocxExtractor,
    HtmlDocumentParser,
    ImageExtractor,
    TextExtractor,
    TiffExtractor,
)
from extractors.image_enhancer import EnhancementProfile, ImageEnhancer
from extractors.ocr_engine import OCREngine, OCRLine, OCRPageResult
from normalizers.date_normalizer import NormalizedDate, extract_dates, normalize_date, normalize_dates_from_text
from normalizers.financial_normalizer import (
    NormalizedFinancial,
    extract_financial_amounts,
    extract_financials,
    format_currency,
    normalize_financial,
)
from normalizers.case_normalizer import (
    NormalizedCaseCitation,
    extract_case_citations,
    extract_case_numbers,
)
from normalizers.entity_normalizer import (
    NormalizedEntity,
    double_metaphone,
    extract_correspondence_parties,
    normalize_entity,
    soundex,
    strip_corporate_suffix,
)


# ============================================================================
# 1. SQLITE SCHEMA DEFINITIONS & VAULT HELPERS
# ============================================================================

VAULT_SCHEMA_SQL = """
PRAGMA foreign_keys = ON;

CREATE TABLE IF NOT EXISTS documents (
    document_id TEXT PRIMARY KEY,
    source_uri TEXT NOT NULL,
    file_name TEXT NOT NULL,
    file_path TEXT NOT NULL,
    file_size_bytes INTEGER NOT NULL CHECK(file_size_bytes >= 0),
    mime_type TEXT NOT NULL,
    file_sha256 TEXT NOT NULL UNIQUE,
    content_sha256 TEXT NOT NULL,
    ingestion_timestamp TEXT NOT NULL,
    document_date TEXT,
    page_count INTEGER NOT NULL DEFAULT 1 CHECK(page_count >= 1),
    extracted_text TEXT,
    ocr_confidence REAL NOT NULL DEFAULT 1.0 CHECK(ocr_confidence >= 0.0 AND ocr_confidence <= 1.0),
    raw_metadata_json TEXT NOT NULL DEFAULT '{}',
    created_at TEXT NOT NULL DEFAULT (strftime('%Y-%m-%dT%H:%M:%SZ', 'now'))
);

CREATE TABLE IF NOT EXISTS entities (
    entity_id TEXT PRIMARY KEY,
    canonical_name TEXT NOT NULL,
    entity_category TEXT NOT NULL CHECK(
        entity_category IN (
            'INDIVIDUAL',
            'MUNICIPAL_BODY',
            'FINANCIAL_INSTITUTION',
            'PROPERTY_MANAGEMENT',
            'LEGAL_AGENCY',
            'COMMERCIAL_ENTITY',
            'OTHER'
        )
    ),
    role_or_title TEXT,
    primary_jurisdiction TEXT,
    aliases_json TEXT NOT NULL DEFAULT '[]',
    metadata_json TEXT NOT NULL DEFAULT '{}',
    created_at TEXT NOT NULL DEFAULT (strftime('%Y-%m-%dT%H:%M:%SZ', 'now')),
    updated_at TEXT NOT NULL DEFAULT (strftime('%Y-%m-%dT%H:%M:%SZ', 'now'))
);

CREATE TABLE IF NOT EXISTS entity_mentions (
    mention_id TEXT PRIMARY KEY,
    document_id TEXT NOT NULL REFERENCES documents(document_id) ON DELETE CASCADE,
    entity_id TEXT NOT NULL REFERENCES entities(entity_id) ON DELETE CASCADE,
    raw_mention_text TEXT NOT NULL,
    char_offset_start INTEGER CHECK(char_offset_start >= 0),
    char_offset_end INTEGER CHECK(char_offset_end >= char_offset_start),
    page_number INTEGER NOT NULL DEFAULT 1 CHECK(page_number >= 1),
    context_snippet TEXT,
    confidence_score REAL NOT NULL DEFAULT 1.0 CHECK(confidence_score >= 0.0 AND confidence_score <= 1.0),
    extraction_method TEXT NOT NULL CHECK(extraction_method IN ('REGEX', 'NER', 'MANUAL', 'HYBRID')),
    created_at TEXT NOT NULL DEFAULT (strftime('%Y-%m-%dT%H:%M:%SZ', 'now'))
);

CREATE TABLE IF NOT EXISTS timeline_events (
    event_id TEXT PRIMARY KEY,
    document_id TEXT REFERENCES documents(document_id) ON DELETE SET NULL,
    event_date_iso TEXT NOT NULL,
    event_year INTEGER NOT NULL,
    event_month INTEGER CHECK(event_month BETWEEN 1 AND 12),
    event_day INTEGER CHECK(event_day BETWEEN 1 AND 31),
    event_type TEXT NOT NULL CHECK(
        event_type IN (
            'JUDICIAL_FILING',
            'REGULATORY_NOTICE',
            'LEGISLATIVE_ACTION',
            'FINANCIAL_TRANSACTION',
            'INCIDENT_LOG',
            'ARREST_SEARCH',
            'RETALIATION_ACTION',
            'ENVIRONMENTAL_HAZARD',
            'OTHER'
        )
    ),
    title TEXT NOT NULL,
    description TEXT NOT NULL,
    raw_snippet TEXT,
    primary_entity_id TEXT REFERENCES entities(entity_id) ON DELETE SET NULL,
    location TEXT,
    jurisdiction TEXT,
    confidence_score REAL NOT NULL DEFAULT 1.0 CHECK(confidence_score >= 0.0 AND confidence_score <= 1.0),
    chronological_rank INTEGER,
    created_at TEXT NOT NULL DEFAULT (strftime('%Y-%m-%dT%H:%M:%SZ', 'now'))
);

CREATE TABLE IF NOT EXISTS financial_transactions (
    transaction_id TEXT PRIMARY KEY,
    document_id TEXT REFERENCES documents(document_id) ON DELETE SET NULL,
    event_id TEXT REFERENCES timeline_events(event_id) ON DELETE SET NULL,
    transaction_date_iso TEXT NOT NULL,
    amount REAL NOT NULL CHECK(amount >= 0.0),
    currency TEXT NOT NULL DEFAULT 'USD',
    sender_entity_id TEXT REFERENCES entities(entity_id) ON DELETE SET NULL,
    recipient_entity_id TEXT REFERENCES entities(entity_id) ON DELETE SET NULL,
    sender_raw_text TEXT,
    recipient_raw_text TEXT,
    payment_method TEXT NOT NULL CHECK(
        payment_method IN ('WIRE', 'CHECK', 'CASH', 'ESCROW', 'GRANT', 'BRIBERY_CONDUIT', 'INVOICE', 'UNKNOWN')
    ),
    account_or_check_num TEXT,
    transaction_purpose TEXT,
    is_predicate_act INTEGER NOT NULL DEFAULT 0 CHECK(is_predicate_act IN (0, 1)),
    raw_snippet TEXT,
    created_at TEXT NOT NULL DEFAULT (strftime('%Y-%m-%dT%H:%M:%SZ', 'now'))
);

CREATE TABLE IF NOT EXISTS relationships (
    relationship_id TEXT PRIMARY KEY,
    source_entity_id TEXT NOT NULL REFERENCES entities(entity_id) ON DELETE CASCADE,
    target_entity_id TEXT NOT NULL REFERENCES entities(entity_id) ON DELETE CASCADE,
    relationship_type TEXT NOT NULL CHECK(
        relationship_type IN (
            'OFFICER_OF',
            'EMPLOYED_BY',
            'CONTROLLED_BY',
            'TRANSFERRED_FUNDS_TO',
            'SUED_BY',
            'REPRESENTED_BY',
            'CO_CONSPIRATOR_WITH',
            'RETALIATED_AGAINST',
            'SUBMITTED_BID_TO',
            'ISSUED_NOTICE_TO',
            'CONNECTED_TO'
        )
    ),
    direction TEXT NOT NULL DEFAULT 'DIRECTED' CHECK(direction IN ('DIRECTED', 'BIDIRECTIONAL')),
    confidence REAL NOT NULL DEFAULT 1.0 CHECK(confidence >= 0.0 AND confidence <= 1.0),
    valid_from TEXT,
    valid_to TEXT,
    source_document_id TEXT REFERENCES documents(document_id) ON DELETE SET NULL,
    evidence_summary TEXT,
    created_at TEXT NOT NULL DEFAULT (strftime('%Y-%m-%dT%H:%M:%SZ', 'now')),
    CHECK(source_entity_id <> target_entity_id)
);

CREATE TABLE IF NOT EXISTS schema_invariants_log (
    audit_id INTEGER PRIMARY KEY AUTOINCREMENT,
    audit_timestamp TEXT NOT NULL DEFAULT (strftime('%Y-%m-%dT%H:%M:%SZ', 'now')),
    tier_level TEXT NOT NULL,
    merkle_root_sha256 TEXT NOT NULL,
    documents_count INTEGER NOT NULL,
    entities_count INTEGER NOT NULL,
    events_count INTEGER NOT NULL,
    transactions_count INTEGER NOT NULL,
    relationships_count INTEGER NOT NULL,
    foreign_key_violations INTEGER NOT NULL DEFAULT 0,
    chronological_inversions INTEGER NOT NULL DEFAULT 0,
    verification_status TEXT NOT NULL CHECK(verification_status IN ('PASSED', 'FAILED'))
);

CREATE INDEX IF NOT EXISTS idx_documents_file_sha256 ON documents(file_sha256);
CREATE INDEX IF NOT EXISTS idx_documents_mime ON documents(mime_type);
CREATE INDEX IF NOT EXISTS idx_entities_canonical_name ON entities(canonical_name);
CREATE INDEX IF NOT EXISTS idx_entities_category ON entities(entity_category);
CREATE INDEX IF NOT EXISTS idx_entity_mentions_doc ON entity_mentions(document_id);
CREATE INDEX IF NOT EXISTS idx_entity_mentions_ent ON entity_mentions(entity_id);
CREATE INDEX IF NOT EXISTS idx_timeline_events_date ON timeline_events(event_date_iso);
CREATE INDEX IF NOT EXISTS idx_timeline_events_entity ON timeline_events(primary_entity_id);
CREATE INDEX IF NOT EXISTS idx_timeline_events_type ON timeline_events(event_type);
CREATE INDEX IF NOT EXISTS idx_financial_trx_date ON financial_transactions(transaction_date_iso);
CREATE INDEX IF NOT EXISTS idx_financial_trx_sender ON financial_transactions(sender_entity_id);
CREATE INDEX IF NOT EXISTS idx_financial_trx_recipient ON financial_transactions(recipient_entity_id);
CREATE INDEX IF NOT EXISTS idx_relationships_source_target ON relationships(source_entity_id, target_entity_id);
CREATE INDEX IF NOT EXISTS idx_relationships_type ON relationships(relationship_type);
"""


@pytest.fixture
def temp_vault_db(tmp_path: Path) -> Generator[Tuple[sqlite3.Connection, Path], None, None]:
    """
    Creates a temporary SQLite database initialized with the complete 7-table schema,
    foreign keys enforced, and WAL mode active.
    """
    db_path = tmp_path / "test_timeline_vault.db"
    conn = sqlite3.connect(str(db_path))
    conn.execute("PRAGMA foreign_keys = ON;")
    conn.execute("PRAGMA journal_mode = WAL;")
    conn.executescript(VAULT_SCHEMA_SQL)
    conn.commit()
    yield conn, db_path
    conn.close()


@pytest.fixture
def in_memory_vault_db() -> Generator[sqlite3.Connection, None, None]:
    """
    Creates an isolated in-memory SQLite database initialized with the 7-table schema.
    """
    conn = sqlite3.connect(":memory:")
    conn.execute("PRAGMA foreign_keys = ON;")
    conn.executescript(VAULT_SCHEMA_SQL)
    conn.commit()
    yield conn
    conn.close()


# ============================================================================
# 2. CRYPTOGRAPHIC & MERKLE TREE FIXTURES
# ============================================================================

def compute_canonical_json_bytes(data: Any) -> bytes:
    """Serializes data according to RFC 8785 JSON Canonicalization Scheme (JCS)."""
    return json.dumps(
        data,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":")
    ).encode("utf-8")


def compute_canonical_json_sha256(data: Any) -> str:
    """Computes SHA-256 hex digest of RFC 8785 canonical JSON bytes."""
    return hashlib.sha256(compute_canonical_json_bytes(data)).hexdigest().lower()


def compute_merkle_root(hashes: List[str]) -> str:
    """
    Computes binary Merkle tree root from an ordered list of SHA-256 hex strings.
    If hashes is empty, returns SHA-256 of empty string.
    """
    if not hashes:
        return hashlib.sha256(b"").hexdigest().lower()
    
    current_level = [h.lower() for h in hashes]
    while len(current_level) > 1:
        next_level = []
        for i in range(0, len(current_level), 2):
            left = current_level[i]
            right = current_level[i + 1] if i + 1 < len(current_level) else left
            combined = (left + right).encode("utf-8")
            parent_hash = hashlib.sha256(combined).hexdigest().lower()
            next_level.append(parent_hash)
        current_level = next_level
    
    return current_level[0]


@pytest.fixture
def merkle_tools() -> Dict[str, Callable]:
    """Exposes Merkle tree and canonical JSON utilities to test cases."""
    return {
        "canonical_json_bytes": compute_canonical_json_bytes,
        "canonical_json_sha256": compute_canonical_json_sha256,
        "merkle_root": compute_merkle_root,
    }


# ============================================================================
# 3. SYNTHETIC DOCUMENT & ARTIFACT GENERATORS
# ============================================================================

@pytest.fixture
def make_synthetic_pdf(tmp_path: Path) -> Callable[..., Path]:
    """
    Factory fixture to generate synthetic multi-page PDFs with text and optional images.
    """
    def _create_pdf(
        filename: str = "sample.pdf",
        pages_content: Optional[List[str]] = None,
        metadata: Optional[Dict[str, str]] = None
    ) -> Path:
        target_path = tmp_path / filename
        target_path.parent.mkdir(parents=True, exist_ok=True)

        pages = pages_content or [
            "UNITED STATES DISTRICT COURT\nCENTRAL DISTRICT OF CALIFORNIA\nCase No. 8:23-cr-00108-CJC\nUNITED STATES v. HARRY SIDHU\nPLEA AGREEMENT\nDate: August 16, 2023\nAmount: $320,000,000"
        ]

        doc = pymupdf.open()
        for idx, page_text in enumerate(pages):
            page = doc.new_page(width=612, height=792)  # Letter size
            # Insert text at top
            page.insert_text((50, 72), page_text, fontsize=11)

        if metadata:
            doc.set_metadata(metadata)

        doc.save(str(target_path))
        doc.close()
        return target_path

    return _create_pdf


@pytest.fixture
def make_synthetic_image(tmp_path: Path) -> Callable[..., Path]:
    """
    Factory fixture to generate synthetic image files with text, skew, and contrast adjustments.
    """
    def _create_image(
        filename: str = "document_scan.png",
        text_lines: Optional[List[str]] = None,
        skew_angle: float = 0.0,
        low_contrast: bool = False,
        dimensions: Tuple[int, int] = (1200, 1600),
    ) -> Path:
        target_path = tmp_path / filename
        target_path.parent.mkdir(parents=True, exist_ok=True)

        bg_color = (200, 200, 200) if low_contrast else (255, 255, 255)
        fg_color = (160, 160, 160) if low_contrast else (0, 0, 0)

        img = Image.new("RGB", dimensions, color=bg_color)
        draw = ImageDraw.Draw(img)

        lines = text_lines or [
            "STATE OF CALIFORNIA - HOUSING AND COMMUNITY DEVELOPMENT",
            "OFFICIAL NOTICE OF VIOLATION",
            "Date: December 8, 2021",
            "Statute: Cal. Gov. Code § 54220",
            "Penalty Assessment: $96,000,000.00",
            "Target: City of Anaheim / Angel Stadium Parcel"
        ]

        y_pos = 100
        for line in lines:
            draw.text((80, y_pos), line, fill=fg_color)
            y_pos += 60

        if abs(skew_angle) > 0.1:
            img = img.rotate(skew_angle, expand=False, fillcolor=bg_color)

        img.save(str(target_path))
        return target_path

    return _create_image


@pytest.fixture
def make_synthetic_docx(tmp_path: Path) -> Callable[..., Path]:
    """
    Factory fixture to generate synthetic Microsoft Word (.docx) documents.
    """
    def _create_docx(
        filename: str = "contract.docx",
        paragraphs: Optional[List[str]] = None,
        table_rows: Optional[List[List[str]]] = None
    ) -> Path:
        target_path = tmp_path / filename
        target_path.parent.mkdir(parents=True, exist_ok=True)

        doc = docx.Document()
        doc.add_heading("FORENSIC INVESTIGATION MEMORANDUM", 0)

        paras = paragraphs or [
            "Case: 30-2021-01201327-CL-UD-CJC",
            "Parties: Woodbridge Meadows Apartments LLC v. DiMarcello",
            "Date: 2021-06-29",
            "Summary: Default Judgment entered for $50,000.00."
        ]
        for p in paras:
            doc.add_paragraph(p)

        if table_rows:
            table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            for r_idx, row in enumerate(table_rows):
                for c_idx, val in enumerate(row):
                    table.cell(r_idx, c_idx).text = val

        doc.save(str(target_path))
        return target_path

    return _create_docx


@pytest.fixture
def make_synthetic_eml(tmp_path: Path) -> Callable[..., Path]:
    """
    Factory fixture to generate RFC 2822 / MIME .eml files with headers and attachments.
    """
    def _create_eml(
        filename: str = "message.eml",
        sender: str = "todd.ament@anaheimchamber.org",
        recipients: List[str] = None,
        subject: str = "Re: Angel Stadium Appraisal Briefing",
        date_str: str = "Tue, 21 May 2019 06:04:00 -0700",
        body_text: str = "Mayor Sidhu agreed to provide the confidential appraisal to SRB Management.",
        attachment_bytes: Optional[bytes] = None,
        attachment_filename: Optional[str] = None,
    ) -> Path:
        target_path = tmp_path / filename
        target_path.parent.mkdir(parents=True, exist_ok=True)

        to_addrs = recipients or ["harry.sidhu@anaheim.net", "jeffrey.flint@fpsstrategies.com"]

        msg = EmailMessage()
        msg["From"] = sender
        msg["To"] = ", ".join(to_addrs)
        msg["Subject"] = subject
        msg["Date"] = date_str
        msg.set_content(body_text)

        if attachment_bytes and attachment_filename:
            msg.add_attachment(
                attachment_bytes,
                maintype="application",
                subtype="pdf",
                filename=attachment_filename
            )

        with open(target_path, "wb") as f:
            f.write(msg.as_bytes())

        return target_path

    return _create_eml


@pytest.fixture
def make_synthetic_archive(tmp_path: Path) -> Callable[..., Path]:
    """
    Factory fixture to generate synthetic ZIP, TAR, or GZ archives containing multiple files.
    """
    def _create_archive(
        archive_name: str = "evidence_package.zip",
        files_dict: Optional[Dict[str, bytes]] = None,
    ) -> Path:
        target_path = tmp_path / archive_name
        target_path.parent.mkdir(parents=True, exist_ok=True)

        files = files_dict or {
            "plea_agreement.txt": b"Case No. 8:23-cr-00108-CJC United States v. Harry Sidhu",
            "invoice_14098.txt": b"Quantum Auto Dismantler Invoice #14098 Total: $1,250.00",
            "subfolder/notice.txt": b"HCD Notice of Violation Cal. Gov. Code 54220 $96M"
        }

        if archive_name.endswith(".zip"):
            with zipfile.ZipFile(target_path, "w", zipfile.ZIP_DEFLATED) as zf:
                for arcname, data in files.items():
                    zf.writestr(arcname, data)
        elif archive_name.endswith((".tar", ".tar.gz", ".tgz")):
            mode = "w:gz" if archive_name.endswith((".tar.gz", ".tgz")) else "w"
            with tarfile.open(target_path, mode) as tf:
                for arcname, data in files.items():
                    ti = tarfile.TarInfo(name=arcname)
                    ti.size = len(data)
                    tf.addfile(ti, io.BytesIO(data))
        else:
            raise ValueError(f"Unsupported archive extension: {archive_name}")

        return target_path

    return _create_archive


# ============================================================================
# 4. INVESTIGATIVE DOMAIN CORPORA (SCENARIOS 1 - 9 DATA)
# ============================================================================

@pytest.fixture
def angel_stadium_records() -> Dict[str, Any]:
    """
    Authoritative test dataset for Scenario 1: Anaheim Angel Stadium Public Corruption.
    """
    return {
        "cases": [
            {
                "docket": "8:23-cr-00108-CJC",
                "court": "USDC CDCA",
                "defendant": "Harry Sidhu",
                "role": "Former Mayor of Anaheim",
                "plea_date": "2023-08-16",
                "statutes": ["18 U.S.C. § 1343", "18 U.S.C. § 1001"],
            },
            {
                "docket": "8:22-cr-00078-CJC",
                "court": "USDC CDCA",
                "defendant": "Todd Ament",
                "role": "Former CEO Anaheim Chamber of Commerce",
                "plea_date": "2022-07-15",
                "statutes": ["18 U.S.C. § 1343", "26 U.S.C. § 7206(1)"],
            },
            {
                "docket": "8:23-cr-00009-CJC",
                "court": "USDC CDCA",
                "defendant": "Melahat Rafiei",
                "role": "Political Consultant / FPS Strategies",
                "plea_date": "2023-01-19",
                "statutes": ["18 U.S.C. § 1343"],
            }
        ],
        "financials": [
            {"raw": "$320,000,000", "float": 320000000.0, "cents": 32000000000, "purpose": "Original Stadium Land Sale Price"},
            {"raw": "$96,000,000", "float": 96000000.0, "cents": 9600000000, "purpose": "HCD Surplus Land Act Penalty (30%)"},
            {"raw": "$50,000,000", "float": 50000000.0, "cents": 5000000000, "purpose": "Refunded Escrow Deposit to SRB Management"},
            {"raw": "$1,000,000", "float": 1000000.0, "cents": 100000000, "purpose": "Sidhu Campaign Solicitations Target"},
        ],
        "timeline_events": [
            {"date": "2021-12-08", "type": "REGULATORY_NOTICE", "title": "HCD Official Notice of Violation issued to Anaheim"},
            {"date": "2022-05-16", "type": "ARREST_SEARCH", "title": "FBI Search Warrant Affidavit unsealed against Mayor Sidhu"},
            {"date": "2022-05-24", "type": "LEGISLATIVE_ACTION", "title": "Anaheim City Council votes unanimously on Res. 2022-064 voiding stadium deal"},
            {"date": "2023-08-16", "type": "JUDICIAL_FILING", "title": "Harry Sidhu signs 4-count federal Plea Agreement in CDCA"},
        ]
    }


@pytest.fixture
def unlawful_detainer_records() -> Dict[str, Any]:
    """
    Authoritative test dataset for Scenario 2: California Superior Court Unlawful Detainer Docket.
    """
    return {
        "case_number": "30-2021-01201327-CL-UD-CJC",
        "jurisdiction": "California Superior Court (Orange County CJC)",
        "plaintiff": "Woodbridge Meadows Apartments LLC",
        "defendant": "Anthony DiMarcello",
        "plaintiff_counsel": "Wallace, Richardson, Sontag & Le LLP",
        "peremptory_judge": "Carmen Luege",
        "total_roa_entries": 61,
        "key_dates": [
            {"date": "2021-05-19", "event": "Unlawful Detainer Complaint Filed (ROA #1)"},
            {"date": "2021-06-29", "event": "Default Judgment 1 Entered (ROA #14)"},
            {"date": "2021-12-22T15:11:00Z", "event": "Judge Luege Chambers Order Granting Ex Parte Stay (ROA #34)"},
            {"date": "2021-12-22T16:29:00Z", "event": "Tactical 4:29 PM Cal. CCP § 170.6 Peremptory Challenge Filed striking Judge Luege (ROA #35)"},
            {"date": "2022-02-04", "event": "Default Judgment 3 Entered (ROA #52)"}
        ]
    }
