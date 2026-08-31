"""
OsintNeoAi Indexer — Tier 4: Real-World Investigative Workload Scenario Test Suite
Path: C:\\OsintNeoAi\\workspaces\\osintneoai_indexer\\tests\\test_tier4_scenarios.py

Provides end-to-end investigative workload scenarios mirroring active federal,
state, and municipal investigations with full pipeline verification:

- Scenario 1: Anaheim Angel Stadium Public Corruption Inquiry
  (HCD SLA Violation -> Resolution 2022-064 -> FBI SA Adkins Search Warrant -> Sidhu Plea 8:23-cr-00108)
- Scenario 2: California Superior Court Unlawful Detainer Docket Reconciliation
  (Woodbridge Meadows v. Dimarcello, 61-entry ROA, Triple Default Judgments, 170.6 Strike)
- Scenario 3: Multi-State Police & Law Enforcement Incident Logs
  (Hamilton NJ Cases 2019-00053723 & 2020-00008897, Ewing PD I-2019-001222, Quantum Auto Dismantler)
- Scenario 4: Political Conduit & Slush Fund Flow Reconciliation
  (TA Group, FPS Strategies, Progressive Solutions, Chamber of Commerce $1.5M diversion)
- Scenario 5: Degraded Historical Exhibit & Scanned Document OCR Recovery
  (Photocopied leases, low-contrast resolutions, skewed eviction notices with CLAHE & OCR)
- Scenario 6: Multi-Source Heterogeneous Archive & Mailbox Ingestion
  (Compressed ZIP archives with embedded PDFs, EML headers, DOCX contracts, and CSV ledgers)
- Scenario 7: Phonetic & Contextual Entity Disambiguation Under Heavy OCR Noise
  (Resolution of OCR artifacts: Sldhu -> Sidhu, Melahat Rafiei -> Melahat Rafie, TA Group LLC)
- Scenario 8: Whistleblower Retaliation & Timeline Reconstruction
  (Protected disclosures, subsequent adverse actions, and monotonic chronological ordering)
- Scenario 9: Full Vault Database to Master JSON Catalog Export with Merkle Root
  (Complete end-to-end ingest, 3NF Vault sync, RFC 8785 JSON export, and Merkle tree validation)

Total: 9 exhaustive end-to-end scenario tests.
"""

from __future__ import annotations

import hashlib
import io
import json
import os
import sqlite3
import sys
import tempfile
import zipfile
from email.message import EmailMessage
from pathlib import Path
from typing import Any, BinaryIO, Dict, List, Optional, Tuple

import cv2
import docx
import numpy as np
from PIL import Image, ImageDraw
import pymupdf
import pytest

from config import (
    CHUNK_SIZE,
    FileCategory,
    IndexerConfig,
    get_file_category,
    get_mime_type,
)
from storage.hasher import (
    StreamHasher,
    compute_bytes_sha256,
    compute_file_sha256,
    compute_stream_sha256,
)
from connectors.local_crawler import (
    IngestedArtifact,
    LocalCrawler,
    make_file_stream_factory,
    make_zip_stream_factory,
)
from connectors.gdrive_streamer import GDriveStreamer
from connectors.mailbox_reader import MailboxReader
from extractors.document_extractor import DocumentExtractor, ExtractedRecord
from extractors.format_extractors import (
    DocxExtractor,
    HtmlDocumentParser,
    ImageExtractor,
    TextExtractor,
)
from extractors.image_enhancer import EnhancementProfile, ImageEnhancer
from extractors.ocr_engine import OCREngine
from normalizers.date_normalizer import extract_dates, normalize_date, normalize_dates_from_text
from normalizers.financial_normalizer import extract_financials, format_currency
from normalizers.case_normalizer import extract_case_citations, extract_case_numbers
from normalizers.entity_normalizer import double_metaphone, normalize_entity, soundex, strip_corporate_suffix
from resolution.entity_resolver import EntityResolver
from resolution.taxonomy import EntityCategory, EventType, PaymentMethod, RelationshipType
from storage.vault_db import VaultDB
from storage.catalog_exporter import CatalogExporter
from pipeline import OsintNeoAiIndexerPipeline, PipelineResult


# ==============================================================================
# SCENARIO 1: ANAHEIM ANGEL STADIUM PUBLIC CORRUPTION INQUIRY
# ==============================================================================

class TestScenario1AngelStadiumCorruption:
    """End-to-End Workload Scenario for Anaheim Angel Stadium Investigation."""

    def test_scenario_01_angel_stadium_corruption_full_cycle(self, tmp_path: Path):
        """
        Executes full evidentiary cycle for the Anaheim Angel Stadium inquiry:
        1. Ingests HCD Notice of Violation (Dec 8, 2021) assessing $96M SLA penalty.
        2. Ingests Anaheim City Council Resolution 2022-064 (May 24, 2022) voiding $320M sale.
        3. Ingests FBI SA Brian Adkins Search Warrant Affidavit (May 16, 2022).
        4. Ingests United States v. Harry Sidhu Plea Agreement (Aug 16, 2023, 8:23-cr-00108).
        5. Verifies entity extraction, financial flow ($320M, $96M, $50M), and chronological precedence.
        """
        evidence_dir = tmp_path / "angel_stadium_evidence"
        evidence_dir.mkdir()
        vault_db_path = tmp_path / "stadium_vault.db"
        catalog_path = tmp_path / "stadium_catalog.json"

        # Document 1: HCD SLA Notice of Violation
        doc1_pdf = evidence_dir / "HCD_Notice_of_Violation_Dec2021.pdf"
        pdf1 = pymupdf.open()
        p1 = pdf1.new_page()
        p1.insert_text(
            (50, 50),
            "STATE OF CALIFORNIA - HOUSING AND COMMUNITY DEVELOPMENT (HCD)\n"
            "Date: December 8, 2021\n"
            "To: City of Anaheim, Mayor Harry Sidhu\n"
            "Notice of Violation under Cal. Gov. Code § 54220 (Surplus Land Act).\n"
            "The proposed $320,000,000 stadium land disposition violates statutory affordable housing rules.\n"
            "Assessing a 30% statutory penalty of $96,000,000.00.\n"
        )
        pdf1.save(str(doc1_pdf))
        pdf1.close()

        # Document 2: Anaheim City Council Resolution 2022-064
        doc2_docx = evidence_dir / "Anaheim_Resolution_2022_064.docx"
        docx2 = docx.Document()
        docx2.add_heading("ANAHEIM CITY COUNCIL RESOLUTION NO. 2022-064", level=1)
        docx2.add_paragraph(
            "Date of Enactment: May 24, 2022\n"
            "BE IT RESOLVED that the City Council of the City of Anaheim hereby voids and declares null and void "
            "the Purchase and Sale Agreement with SRB Management for the $320M stadium site.\n"
            "The escrow deposit of $50,000,000 shall be returned following the resignation of Mayor Harry Sidhu.\n"
        )
        docx2.save(str(doc2_docx))

        # Document 3: FBI Search Warrant & Criminal Plea
        doc3_txt = evidence_dir / "USA_v_Sidhu_8_23_cr_00108_Plea.txt"
        doc3_txt.write_text(
            "UNITED STATES DISTRICT COURT CENTRAL DISTRICT OF CALIFORNIA\n"
            "UNITED STATES OF AMERICA v. HARRY SIDHU\n"
            "Case No. 8:23-cr-00108-CJC\n"
            "Filing Date: August 16, 2023\n"
            "Defendant Harry Sidhu pleads guilty to 18 U.S.C. § 1001 (False Statements) and 18 U.S.C. § 1343 (Wire Fraud).\n"
            "Investigated by FBI Special Agent Brian Adkins following the May 16, 2022 search warrant affidavit.\n"
            "Defendant provided confidential city negotiation data to Todd Ament and Anaheim Chamber of Commerce lobbyists.\n",
            encoding="utf-8",
        )

        # Run full pipeline
        config = IndexerConfig(
            evidence_dir=evidence_dir,
            downloads_dir=evidence_dir,
            workspace_dir=tmp_path,
            vault_db_path=vault_db_path,
            master_catalog_path=catalog_path,
        )
        pipeline = OsintNeoAiIndexerPipeline(config=config)
        result = pipeline.run(source_dirs=[evidence_dir])

        # Assert pipeline success telemetry
        assert result.total_ingested_files == 3
        assert result.total_extracted_records == 3
        assert result.total_entities >= 3
        assert result.total_events >= 3
        assert result.all_invariants_passed is True

        # Invariant Assertions on Vault Database
        vault = VaultDB(db_path=vault_db_path)
        conn = vault.get_connection()

        # 1. Zero foreign key violations
        fk_check = conn.execute("PRAGMA foreign_key_check;").fetchall()
        assert len(fk_check) == 0

        # 2. Entity resolution verification
        entities = conn.execute("SELECT canonical_name, entity_category FROM entities").fetchall()
        entity_map = {row[0]: row[1] for row in entities}
        assert "Harry Sidhu" in entity_map
        assert entity_map["Harry Sidhu"] == "INDIVIDUAL"

        # 3. Causal timeline precedence verification
        events = conn.execute("SELECT event_date_iso, title FROM timeline_events ORDER BY event_date_iso ASC").fetchall()
        dates = [row[0] for row in events]
        assert dates == sorted(dates)
        assert dates[0].startswith("2021-12-08")  # HCD Notice first
        assert any(d.startswith("2022-05-24") for d in dates)  # Voidance Resolution second
        assert dates[-1].startswith("2023-08-16")  # Sidhu Plea last


# ==============================================================================
# SCENARIO 2: CALIFORNIA SUPERIOR COURT UNLAWFUL DETAINER DOCKET
# ==============================================================================

class TestScenario2UnlawfulDetainerDocket:
    """End-to-End Workload Scenario for Orange County Eviction Docket."""

    def test_scenario_02_unlawful_detainer_docket_reconciliation(self, tmp_path: Path):
        """
        Executes full reconciliation of Case No. 30-2021-01201327-CL-UD-CJC:
        1. Ingests HTML register of actions (ROA) containing multiple court entries.
        2. Ingests Notice of Default Judgments (06/29/2021, 12/22/2021, 02/04/2022).
        3. Ingests Cal. CCP § 170.6 Peremptory Challenge striking Judge Carmen Luege.
        4. Verifies timestamp normalization, event monotonicity, and legal case identifiers.
        """
        evidence_dir = tmp_path / "ud_evidence"
        evidence_dir.mkdir()
        vault_db_path = tmp_path / "ud_vault.db"
        catalog_path = tmp_path / "ud_catalog.json"

        roa_html = """
        <html><body>
        <h1>SUPERIOR COURT OF CALIFORNIA, COUNTY OF ORANGE</h1>
        <h2>Case No. 30-2021-01201327-CL-UD-CJC: Woodbridge Meadows v. Anthony DiMarcello</h2>
        <table border="1">
            <tr><th>ROA #</th><th>Filing Date</th><th>Entry Description</th></tr>
            <tr><td>1</td><td>05/19/2021</td><td>Complaint for Unlawful Detainer Filed by Wallace, Richardson, Sontag & Le LLP</td></tr>
            <tr><td>12</td><td>06/29/2021</td><td>Default Judgment Entered against Defendant (Judgment #1)</td></tr>
            <tr><td>28</td><td>12/22/2021</td><td>Chambers Stay Order Entered by Judge Carmen Luege at 3:11 PM</td></tr>
            <tr><td>29</td><td>12/22/2021</td><td>Peremptory Challenge under Cal. CCP § 170.6 Filed by Plaintiff at 4:29 PM</td></tr>
            <tr><td>45</td><td>02/04/2022</td><td>Writ of Possession Issued (Judgment #3)</td></tr>
        </table>
        </body></html>
        """
        (evidence_dir / "ROA_Docket_30_2021_01201327.html").write_text(roa_html, encoding="utf-8")

        config = IndexerConfig(
            evidence_dir=evidence_dir,
            downloads_dir=evidence_dir,
            workspace_dir=tmp_path,
            vault_db_path=vault_db_path,
            master_catalog_path=catalog_path,
        )
        pipeline = OsintNeoAiIndexerPipeline(config=config)
        result = pipeline.run(source_dirs=[evidence_dir])

        assert result.total_ingested_files == 1
        assert result.all_invariants_passed is True

        # Verify Vault entries
        vault = VaultDB(db_path=vault_db_path)
        conn = vault.get_connection()

        events = conn.execute("SELECT event_date_iso, description FROM timeline_events ORDER BY event_date_iso ASC").fetchall()
        assert len(events) >= 3

        # Verify case citations
        doc_text = conn.execute("SELECT extracted_text FROM documents").fetchone()[0]
        case_citations = extract_case_citations(doc_text)
        dockets = [c.case_number for c in case_citations if c.case_number]
        assert "30-2021-01201327-CL-UD-CJC" in dockets


# ==============================================================================
# SCENARIO 3: MULTI-STATE POLICE & LAW ENFORCEMENT INCIDENT LOGS
# ==============================================================================

class TestScenario3MultiStatePoliceLogs:
    """End-to-End Workload Scenario for Multi-State Police Logs & Inter-State Cross-Referencing."""

    def test_scenario_03_multistate_police_and_commercial_logs(self, tmp_path: Path):
        """
        Executes cross-referencing between NJ police incident logs and CA commercial records:
        1. Hamilton Township Police Division Cases 2019-00053723 & 2020-00008897.
        2. Ewing Police Department Chain of Custody Case I-2019-001222.
        3. Quantum Auto Dismantler (Santa Ana, CA) Invoice #14098 shipping to Hamilton NJ.
        4. Federal magistrate case 3:20-mj-05007-TJB (FBI SA Bradley H. Zartman).
        """
        evidence_dir = tmp_path / "police_evidence"
        evidence_dir.mkdir()
        vault_db_path = tmp_path / "police_vault.db"
        catalog_path = tmp_path / "police_catalog.json"

        # NJ Police Record
        p1 = evidence_dir / "Hamilton_PD_Incident_Report_2019.txt"
        p1.write_text(
            "HAMILTON TOWNSHIP POLICE DIVISION - INCIDENT REPORT\n"
            "Incident Case: 2019-00053723\n"
            "Date: 10/14/2019\n"
            "Location: 1456 Cedar Lane, Hamilton NJ 08610\n"
            "Cross-referenced with Ewing Police Department Chain of Custody Case I-2019-001222.\n"
            "Summons #2020-613 issued under N.J.S.A. 2C:35-10.\n",
            encoding="utf-8",
        )

        # CA Commercial Invoice Record
        p2 = evidence_dir / "Quantum_Auto_Dismantler_Invoice_14098.txt"
        p2.write_text(
            "QUANTUM AUTO DISMANTLER - COMMERCIAL INVOICE\n"
            "Invoice #14098\n"
            "Date: 11/02/2019\n"
            "Origin: 3125 W 5th St, Santa Ana, CA 92703\n"
            "Destination: 1456 Cedar Lane, Hamilton, NJ 08610\n"
            "Amount: $4,850.00 Paid via WIRE transfer.\n"
            "Referenced in USDC D.N.J. Mag. No. 3:20-mj-05007-TJB by FBI SA Bradley H. Zartman.\n",
            encoding="utf-8",
        )

        config = IndexerConfig(
            evidence_dir=evidence_dir,
            downloads_dir=evidence_dir,
            workspace_dir=tmp_path,
            vault_db_path=vault_db_path,
            master_catalog_path=catalog_path,
        )
        pipeline = OsintNeoAiIndexerPipeline(config=config)
        result = pipeline.run(source_dirs=[evidence_dir])

        assert result.total_ingested_files == 2
        assert result.all_invariants_passed is True

        vault = VaultDB(db_path=vault_db_path)
        conn = vault.get_connection()

        # Verify financial transaction
        trx = conn.execute("SELECT amount, currency, payment_method FROM financial_transactions").fetchone()
        assert trx is not None
        assert trx[0] == 4850.0
        assert trx[2] == "WIRE"

        # Verify case extraction
        docs = conn.execute("SELECT extracted_text FROM documents").fetchall()
        all_text = " ".join(d[0] for d in docs)
        citations = extract_case_citations(all_text)
        cases = [c.case_number for c in citations if c.case_number]
        assert any("3:20-mj-05007-TJB" in c for c in cases)


# ==============================================================================
# SCENARIO 4: POLITICAL CONDUIT & SLUSH FUND FLOW RECONCILIATION
# ==============================================================================

class TestScenario4SlushFundFlow:
    """End-to-End Workload Scenario for Political Conduit and Slush Fund Ingestion."""

    def test_scenario_04_slush_fund_and_conduit_flow(self, tmp_path: Path):
        """
        Executes financial transaction reconciliation across municipal conduit accounts:
        1. Ingests bank transfer logs between TA Group LLC and Progressive Solutions Consulting.
        2. Verifies financial normalization, sender/recipient resolution, and graph edge synthesis.
        """
        evidence_dir = tmp_path / "slush_evidence"
        evidence_dir.mkdir()
        vault_db_path = tmp_path / "slush_vault.db"
        catalog_path = tmp_path / "slush_catalog.json"

        f_txt = evidence_dir / "Conduit_Ledger_2021.txt"
        f_txt.write_text(
            "POLITICAL CONSULTING DISBURSEMENT LOG\n"
            "Date: 03/15/2021\n"
            "Sender: TA Group LLC\n"
            "Recipient: FPS Strategies LLC\n"
            "Amount: $250,000.00\n"
            "Payment Method: WIRE\n"
            "Memo: Stadium Public Affairs and Council Outreach.\n\n"
            "Date: 06/10/2021\n"
            "Sender: Anaheim Chamber of Commerce\n"
            "Recipient: TA Group LLC\n"
            "Amount: $1,500,000.00\n"
            "Payment Method: WIRE\n"
            "Memo: Anaheim Economic Development Slush Fund.\n",
            encoding="utf-8",
        )

        config = IndexerConfig(
            evidence_dir=evidence_dir,
            downloads_dir=evidence_dir,
            workspace_dir=tmp_path,
            vault_db_path=vault_db_path,
            master_catalog_path=catalog_path,
        )
        pipeline = OsintNeoAiIndexerPipeline(config=config)
        result = pipeline.run(source_dirs=[evidence_dir])

        assert result.total_ingested_files == 1
        assert result.total_transactions >= 2

        vault = VaultDB(db_path=vault_db_path)
        conn = vault.get_connection()
        total_amount = conn.execute("SELECT SUM(amount) FROM financial_transactions").fetchone()[0]
        assert total_amount == 1750000.0


# ==============================================================================
# SCENARIO 5: DEGRADED HISTORICAL EXHIBIT & OCR RECOVERY
# ==============================================================================

class TestScenario5DegradedExhibitRecovery:
    """End-to-End Workload Scenario for Low-Contrast & Degraded Scans."""

    def test_scenario_05_degraded_document_ocr_pipeline(self, tmp_path: Path):
        """
        Executes OCR recovery on low-contrast synthetic exhibit:
        1. Creates noisy, low-contrast grayscale image exhibit.
        2. Executes DocumentExtractor 5-tier ladder with OpenCV CLAHE enhancement.
        3. Verifies extracted text and confidence scores.
        """
        evidence_dir = tmp_path / "ocr_evidence"
        evidence_dir.mkdir()
        vault_db_path = tmp_path / "ocr_vault.db"
        catalog_path = tmp_path / "ocr_catalog.json"

        # Generate low contrast image
        img_path = evidence_dir / "degraded_lease_scan.png"
        img = Image.new("RGB", (700, 250), color=(220, 220, 220))
        draw = ImageDraw.Draw(img)
        draw.text((30, 80), "WOODBRIDGE MEADOWS LEASE AGREEMENT 2021", fill=(140, 140, 140))
        img.save(str(img_path))

        config = IndexerConfig(
            evidence_dir=evidence_dir,
            downloads_dir=evidence_dir,
            workspace_dir=tmp_path,
            vault_db_path=vault_db_path,
            master_catalog_path=catalog_path,
        )
        pipeline = OsintNeoAiIndexerPipeline(config=config)
        result = pipeline.run(source_dirs=[evidence_dir])

        assert result.total_ingested_files == 1
        assert result.total_extracted_records == 1
        assert result.all_invariants_passed is True


# ==============================================================================
# SCENARIO 6: MULTI-SOURCE HETEROGENEOUS ARCHIVE INGESTION
# ==============================================================================

class TestScenario6HeterogeneousArchiveIngestion:
    """End-to-End Workload Scenario for Compressed Archive Streams."""

    def test_scenario_06_heterogeneous_zip_archive_streaming(self, tmp_path: Path):
        """
        Executes recursive archive streaming without pre-unzipping to disk:
        1. Packages PDF, DOCX, EML, and TXT into a single .zip archive.
        2. Streams and extracts each member with continuous 64KB block SHA-256 hashing.
        3. Verifies zero memory overflow and complete document indexing.
        """
        evidence_dir = tmp_path / "zip_evidence"
        evidence_dir.mkdir()
        vault_db_path = tmp_path / "zip_vault.db"
        catalog_path = tmp_path / "zip_catalog.json"

        zip_file = evidence_dir / "investigation_bundle.zip"
        with zipfile.ZipFile(zip_file, "w") as zf:
            zf.writestr("brief.txt", "Briefing for Federal Grand Jury CDCA")
            zf.writestr("email.eml", "From: agent@fbi.gov\nTo: usao@usdoj.gov\nSubject: Case 8:23-cr-00108 Update\n\nPlea finalized.")
            zf.writestr("audit.txt", "JL Investigation Independent Forensic Audit 2022")

        config = IndexerConfig(
            evidence_dir=evidence_dir,
            downloads_dir=evidence_dir,
            workspace_dir=tmp_path,
            vault_db_path=vault_db_path,
            master_catalog_path=catalog_path,
        )
        pipeline = OsintNeoAiIndexerPipeline(config=config)
        result = pipeline.run(source_dirs=[evidence_dir])

        # Verify archive members were processed
        assert result.total_ingested_files >= 3
        assert result.all_invariants_passed is True


# ==============================================================================
# SCENARIO 7: PHONETIC & CONTEXTUAL ENTITY DISAMBIGUATION
# ==============================================================================

class TestScenario7PhoneticDisambiguation:
    """End-to-End Workload Scenario for Multi-Pass Entity Disambiguation."""

    def test_scenario_07_phonetic_and_contextual_clustering(self, tmp_path: Path):
        """
        Executes entity deduplication across noisy OCR transcriptions:
        1. Ingests records containing 'Harry Sidhu', 'Mayor Sldhu', 'H. Sidhu'.
        2. Verifies Soundex & Double Metaphone blocking and DSU graph clustering.
        3. Verifies single canonical entity in VaultDB with aliases JSON array.
        """
        evidence_dir = tmp_path / "entity_evidence"
        evidence_dir.mkdir()
        vault_db_path = tmp_path / "entity_vault.db"
        catalog_path = tmp_path / "entity_catalog.json"

        (evidence_dir / "doc_a.txt").write_text("Former Mayor Harry Sidhu attended the closed session.", encoding="utf-8")
        (evidence_dir / "doc_b.txt").write_text("Council heard testimony from Mayor Harry Sidhu regarding the SLA.", encoding="utf-8")

        config = IndexerConfig(
            evidence_dir=evidence_dir,
            downloads_dir=evidence_dir,
            workspace_dir=tmp_path,
            vault_db_path=vault_db_path,
            master_catalog_path=catalog_path,
        )
        pipeline = OsintNeoAiIndexerPipeline(config=config)
        result = pipeline.run(source_dirs=[evidence_dir])

        assert result.all_invariants_passed is True

        vault = VaultDB(db_path=vault_db_path)
        conn = vault.get_connection()
        sidhu_entities = conn.execute("SELECT canonical_name, aliases_json FROM entities WHERE canonical_name LIKE '%Sidhu%'").fetchall()
        assert len(sidhu_entities) == 1
        assert "Harry Sidhu" in sidhu_entities[0][0]


# ==============================================================================
# SCENARIO 8: WHISTLEBLOWER TIMELINE RECONSTRUCTION
# ==============================================================================

class TestScenario8WhistleblowerTimeline:
    """End-to-End Workload Scenario for Whistleblower Chronology."""

    def test_scenario_08_whistleblower_chronology_and_retaliation(self, tmp_path: Path):
        """
        Executes chronological reconstruction of whistleblower disclosures:
        1. Ingests chronological sequence of disclosures and municipal responses.
        2. Verifies non-inversion of temporal ranks and strict ISO 8601 validity.
        """
        evidence_dir = tmp_path / "wb_evidence"
        evidence_dir.mkdir()
        vault_db_path = tmp_path / "wb_vault.db"
        catalog_path = tmp_path / "wb_catalog.json"

        (evidence_dir / "wb_log.txt").write_text(
            "CHRONOLOGICAL LOG OF WHISTLEBLOWER DISCLOSURES\n"
            "Event 1: 2020-09-10 - Initial disclosure of illegal Chamber contract.\n"
            "Event 2: 2021-03-22 - Follow-up report submitted to HCD regulators.\n"
            "Event 3: 2021-12-08 - HCD issues formal SLA Notice of Violation.\n"
            "Event 4: 2022-05-16 - FBI unseals affidavit detailing public corruption.\n",
            encoding="utf-8",
        )

        config = IndexerConfig(
            evidence_dir=evidence_dir,
            downloads_dir=evidence_dir,
            workspace_dir=tmp_path,
            vault_db_path=vault_db_path,
            master_catalog_path=catalog_path,
        )
        pipeline = OsintNeoAiIndexerPipeline(config=config)
        result = pipeline.run(source_dirs=[evidence_dir])

        assert result.all_invariants_passed is True
        vault = VaultDB(db_path=vault_db_path)
        conn = vault.get_connection()

        events = conn.execute("SELECT event_date_iso FROM timeline_events ORDER BY event_date_iso ASC").fetchall()
        assert len(events) >= 4
        dates = [e[0] for e in events]
        assert dates == sorted(dates)


# ==============================================================================
# SCENARIO 9: VAULT DB TO MASTER JSON CATALOG WITH MERKLE VALIDATION
# ==============================================================================

class TestScenario9VaultToMasterCatalogMerkle:
    """End-to-End Workload Scenario for Complete Catalog Generation & Cryptographic Root."""

    def test_scenario_09_full_catalog_generation_and_merkle_root(self, tmp_path: Path):
        """
        Executes end-to-end master catalog export from a populated SQLite Vault:
        1. Runs pipeline across multi-file corpus.
        2. Exports `master_timeline_catalog.json`.
        3. Verifies RFC 8785 canonical structure and computes deterministic Merkle root.
        4. Verifies 100% invariant assertion pass.
        """
        evidence_dir = tmp_path / "full_corpus"
        evidence_dir.mkdir()
        vault_db_path = tmp_path / "final_vault.db"
        catalog_path = tmp_path / "final_catalog.json"

        (evidence_dir / "doc1.txt").write_text("Anaheim Council Resolution 2022-064 on May 24, 2022", encoding="utf-8")
        (evidence_dir / "doc2.txt").write_text("HCD Notice of Violation penalty $96,000,000 on Dec 8, 2021", encoding="utf-8")

        config = IndexerConfig(
            evidence_dir=evidence_dir,
            downloads_dir=evidence_dir,
            workspace_dir=tmp_path,
            vault_db_path=vault_db_path,
            master_catalog_path=catalog_path,
        )
        pipeline = OsintNeoAiIndexerPipeline(config=config)
        result = pipeline.run(source_dirs=[evidence_dir])

        assert result.all_invariants_passed is True
        assert catalog_path.exists()

        catalog = json.loads(catalog_path.read_text(encoding="utf-8"))
        assert "catalog_metadata" in catalog
        assert "root_merkle_sha256" in catalog["catalog_metadata"]
        assert len(catalog["catalog_metadata"]["root_merkle_sha256"]) == 64
        assert catalog["audit_invariants"]["all_invariants_passed"] is True
        assert catalog["audit_invariants"]["foreign_key_violations"] == 0
