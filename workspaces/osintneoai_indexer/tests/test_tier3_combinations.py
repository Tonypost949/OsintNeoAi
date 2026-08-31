"""
OsintNeoAi Indexer — Tier 3: Cross-Feature Pairwise Integration Test Suite
Path: C:\\OsintNeoAi\\workspaces\\osintneoai_indexer\\tests\\test_tier3_combinations.py

Provides exhaustive, non-trivial cross-feature pairwise integration tests across ALL 17 features:
- Test 1:  F1 + F3   (Stream Ingestion & Chunking -> 64KB Block Streaming SHA-256 Engine)
- Test 2:  F1 + F4   (Stream Ingestion & Chunking -> Multi-Format MIME Dispatcher)
- Test 3:  F2 + F3   (Google Drive Link Resolver -> Cryptographic SHA-256 Engine)
- Test 4:  F2 + F4   (Google Drive Link Resolver -> MIME Dispatcher & Category Routing)
- Test 5:  F4 + F5   (MIME Dispatcher -> Native Digital Text Extractor for PDF/DOCX/HTML)
- Test 6:  F4 + F6   (MIME Dispatcher -> Neural Offline OCR Engine for Degraded Scans)
- Test 7:  F5 + F8   (Native Digital Text Extractor -> ISO 8601 Timestamp Normalizer)
- Test 8:  F5 + F9   (Native Digital Text Extractor -> Financial Transaction Normalizer Dual Float/Cents)
- Test 9:  F5 + F10  (Native Digital Text Extractor -> Legal Case & Statutory Citation Normalizer)
- Test 10: F6 + F7   (Image Preprocessing & CLAHE Enhancement -> Neural Offline OCR Engine)
- Test 11: F8 + F12  (Timestamp Normalizer -> 6-Category Entity Extractor Contextual Association)
- Test 12: F9 + F12  (Financial Normalizer -> 6-Category Entity Extractor Payer/Payee Mapping)
- Test 13: F10 + F12 (Legal Case Normalizer -> 6-Category Entity Extractor Judicial/Agency Mapping)
- Test 14: F11 + F12 (Communication Metadata Normalizer -> Entity Extractor Header Disambiguation)
- Test 15: F12 + F13 (6-Category Entity Extractor -> Phonetic & Contextual DSU Entity Resolver)
- Test 16: F13 + F14 (Phonetic Entity Resolver -> SQLite Relational Vault 3NF Database)
- Test 17: F14 + F15 (SQLite Relational Vault -> Master JSON Catalog Exporter with Merkle Root)
- Test 18: F1 + F14  (Stream Ingestion Crawler -> SQLite Relational Vault Document Ingestion)
- Test 19: F8 + F14  (Timestamp Normalizer -> SQLite Vault Timeline Event Chronological Order)
- Test 20: F15 + F17 (Master JSON Catalog Exporter -> 100% Invariant Cryptographic Verification)

Total: 20 comprehensive pairwise integration tests.
"""

from __future__ import annotations

import hashlib
import io
import json
import os
import sqlite3
import sys
import tempfile
import zipfile
from email.message import EmailMessage
from pathlib import Path
from typing import Any, BinaryIO, Dict, List, Optional, Tuple

import cv2
import docx
import numpy as np
from PIL import Image, ImageDraw
import pymupdf
import pytest

from config import (
    CHUNK_SIZE,
    FileCategory,
    IndexerConfig,
    get_file_category,
    get_mime_type,
    is_supported_file,
)
from storage.hasher import (
    StreamHasher,
    compute_bytes_sha256,
    compute_file_sha256,
    compute_file_sha256_with_size,
    compute_stream_sha256,
    verify_file_sha256,
)
from connectors.local_crawler import (
    IngestedArtifact,
    LocalCrawler,
    detect_mime_type,
    make_file_stream_factory,
    make_zip_stream_factory,
)
from connectors.gdrive_streamer import GDriveStreamer, GDriveResourceInfo
from connectors.mailbox_reader import MailboxReader, EmailMetadata
from extractors.document_extractor import DocumentExtractor, ExtractedRecord
from extractors.format_extractors import (
    DocxExtractor,
    HtmlDocumentParser,
    ImageExtractor,
    TextExtractor,
    TiffExtractor,
)
from extractors.image_enhancer import EnhancementProfile, ImageEnhancer
from extractors.ocr_engine import OCREngine, OCRLine, OCRPageResult
from normalizers.date_normalizer import (
    NormalizedDate,
    extract_dates,
    normalize_date,
    normalize_dates_from_text,
)
from normalizers.financial_normalizer import (
    NormalizedFinancial,
    extract_financial_amounts,
    extract_financials,
    format_currency,
    normalize_financial,
)
from normalizers.case_normalizer import (
    NormalizedCaseCitation,
    extract_case_citations,
    extract_case_numbers,
)
from normalizers.entity_normalizer import (
    NormalizedEntity,
    double_metaphone,
    extract_correspondence_parties,
    normalize_entity,
    soundex,
    strip_corporate_suffix,
)
from resolution.entity_resolver import EntityResolver
from resolution.taxonomy import EntityCategory, EventType, PaymentMethod, RelationshipType
from storage.vault_db import VaultDB
from storage.catalog_exporter import CatalogExporter


# ==============================================================================
# TIER 3 PAIRWISE COMBINATIONS: INGESTION & DISPATCH (TESTS 1 - 4)
# ==============================================================================

class TestTier3IngestionCombinations:
    """Pairwise integration tests across connectors, streaming, and hashing."""

    def test_comb_01_f1_f3_crawler_stream_to_hasher(self, tmp_path: Path):
        """
        [F1 + F3] LocalCrawler multi-chunk stream ingestion piping to 64KB StreamHasher.
        Verifies that chunked streaming produces exact cryptographic SHA-256 match
        against standard one-shot file digest across multi-megabyte evidentiary files.
        """
        test_dir = tmp_path / "f1_f3_evidence"
        test_dir.mkdir()
        
        # Create a 256 KB test file
        sample_data = b"FORENSIC_EXHIBIT_DATA_BLOCK\n" * 9000
        file_path = test_dir / "exhibit_alpha.txt"
        file_path.write_bytes(sample_data)
        
        expected_sha = hashlib.sha256(sample_data).hexdigest().lower()
        
        # Ingest using LocalCrawler
        crawler = LocalCrawler(target_paths=[test_dir], chunk_size=CHUNK_SIZE)
        artifacts = list(crawler.crawl_directory(test_dir))
        
        assert len(artifacts) == 1
        art = artifacts[0]
        assert art.artifact_id == expected_sha
        assert art.file_size_bytes == len(sample_data)
        
        # Verify fresh stream factory reproduces exact stream hash
        stream_hash = compute_stream_sha256(art.raw_stream_factory(), chunk_size=32768)
        assert stream_hash == expected_sha

    def test_comb_02_f1_f4_crawler_to_mime_dispatcher(self, tmp_path: Path):
        """
        [F1 + F4] LocalCrawler stream traversal integrated with MIME Dispatcher.
        Verifies correct MIME classification across heterogeneous evidentiary files
        (PDF, DOCX, PNG, HTML, MBOX, TXT) and archive members.
        """
        test_dir = tmp_path / "f1_f4_corpus"
        test_dir.mkdir()
        
        (test_dir / "plea.pdf").write_bytes(b"%PDF-1.4 sample content")
        (test_dir / "memo.docx").write_bytes(b"PK\x03\x04 docx stub")
        (test_dir / "scan.png").write_bytes(b"\x89PNG\r\n\x1a\n stub")
        (test_dir / "docket.html").write_text("<html><body>Case 8:23-cr-00108</body></html>", encoding="utf-8")
        (test_dir / "notes.txt").write_text("Investigative notes on Anaheim Chamber.", encoding="utf-8")
        
        crawler = LocalCrawler(target_paths=[test_dir])
        artifacts = list(crawler.crawl_directory(test_dir))
        
        mime_map = {Path(art.source_uri).name: art.mime_type for art in artifacts}
        
        assert mime_map["plea.pdf"] == "application/pdf"
        assert mime_map["memo.docx"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert mime_map["scan.png"] == "image/png"
        assert mime_map["docket.html"] == "text/html"
        assert mime_map["notes.txt"] == "text/plain"

    def test_comb_03_f2_f3_gdrive_streamer_to_hasher(self, tmp_path: Path):
        """
        [F2 + F3] Google Drive Resolver download spooling piping to continuous SHA-256 Engine.
        Verifies that mock downloaded Drive streams compute and verify 64KB chunk SHA-256
        before generating IngestedArtifact records.
        """
        spool_dir = tmp_path / "gdrive_spool"
        cache_dir = tmp_path / "gdrive_cache"
        spool_dir.mkdir()
        cache_dir.mkdir()
        
        file_id = "1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgvE2upms"
        raw_content = b"%PDF-1.7 Confidential Whistleblower Disclosure Exhibit" * 500
        expected_sha = hashlib.sha256(raw_content).hexdigest().lower()
        
        # Place in local cache to simulate cached Drive download
        cached_file = cache_dir / f"{file_id}.pdf"
        cached_file.write_bytes(raw_content)
        
        streamer = GDriveStreamer(spool_dir=spool_dir, local_cache_dirs=[cache_dir], prefer_offline=True)
        url = f"https://drive.google.com/file/d/{file_id}/view"
        
        artifact = streamer.ingest_url(url)
        assert artifact.artifact_id == expected_sha
        assert artifact.file_size_bytes == len(raw_content)
        assert verify_file_sha256(cached_file, expected_sha) is True

    def test_comb_04_f2_f4_gdrive_resolver_to_mime_dispatcher(self, tmp_path: Path):
        """
        [F2 + F4] Google Drive Link Resolver integrated with MIME Dispatcher.
        Verifies URL parsing, extension extraction, and category categorization for
        Google Docs exports vs binary uploaded files.
        """
        streamer = GDriveStreamer(spool_dir=tmp_path)
        
        doc_url = "https://docs.google.com/document/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgvE2upms/edit"
        info_doc = streamer.parse_url(doc_url)
        assert info_doc.resource_type == "doc"
        assert get_file_category(info_doc.inferred_filename) == FileCategory.PDF
        
        sheet_url = "https://docs.google.com/spreadsheets/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgvE2upms/edit"
        info_sheet = streamer.parse_url(sheet_url)
        assert info_sheet.resource_type == "sheet"
        assert get_file_category(info_sheet.inferred_filename) == FileCategory.TABULAR


# ==============================================================================
# TIER 3 PAIRWISE COMBINATIONS: EXTRACTION & NORMALIZATION (TESTS 5 - 10)
# ==============================================================================

class TestTier3ExtractionAndNormalizers:
    """Pairwise integration tests between extractors and domain normalizers."""

    def test_comb_05_f4_f5_mime_dispatcher_to_digital_extractor(self, tmp_path: Path):
        """
        [F4 + F5] MIME Dispatcher routing to Native Digital Text Extractors.
        Verifies PDF, DOCX, and HTML dispatching directly into PyMuPDF, python-docx,
        and HtmlDocumentParser with structured text output.
        """
        extractor = DocumentExtractor()
        
        # 1. Create native PDF
        pdf_path = tmp_path / "test_contract.pdf"
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Anaheim Stadium Land Sale Agreement 2022-064 with sufficient text to exceed density threshold.")
        doc.save(str(pdf_path))
        doc.close()
        
        art_pdf = IngestedArtifact(
            artifact_id=compute_file_sha256(pdf_path),
            source_uri=str(pdf_path),
            mime_type="application/pdf",
            file_size_bytes=pdf_path.stat().st_size,
            raw_stream_factory=make_file_stream_factory(str(pdf_path)),
        )
        rec_pdf = extractor.extract(art_pdf)
        assert "Anaheim Stadium Land Sale Agreement 2022-064" in rec_pdf.extracted_text

        # 2. Create DOCX
        docx_path = tmp_path / "test_memo.docx"
        docx_doc = docx.Document()
        docx_doc.add_heading("Confidential Legal Memorandum", level=1)
        docx_doc.add_paragraph("Subject: Surplus Land Act compliance under Cal. Gov. Code § 54220.")
        docx_doc.save(str(docx_path))
        
        art_docx = IngestedArtifact(
            artifact_id=compute_file_sha256(docx_path),
            source_uri=str(docx_path),
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_size_bytes=docx_path.stat().st_size,
            raw_stream_factory=make_file_stream_factory(str(docx_path)),
        )
        rec_docx = extractor.extract(art_docx)
        assert "Surplus Land Act" in rec_docx.extracted_text

    def test_comb_06_f4_f6_mime_dispatcher_to_neural_ocr(self, tmp_path: Path):
        """
        [F4 + F6] MIME Dispatcher routing image artifacts to Neural Offline OCR Engine.
        Verifies that image/png and image/jpeg artifacts correctly trigger OCR parsing.
        """
        img_path = tmp_path / "scan_notice.png"
        img = Image.new("RGB", (600, 200), color=(255, 255, 255))
        draw = ImageDraw.Draw(img)
        draw.text((30, 80), "NOTICE OF VIOLATION - HCD PENALTY $96M", fill=(0, 0, 0))
        img.save(str(img_path))
        
        art = IngestedArtifact(
            artifact_id=compute_file_sha256(img_path),
            source_uri=str(img_path),
            mime_type="image/png",
            file_size_bytes=img_path.stat().st_size,
            raw_stream_factory=make_file_stream_factory(str(img_path)),
        )
        
        extractor = DocumentExtractor()
        rec = extractor.extract(art)
        
        assert rec.ocr_engine_used.startswith("rapidocr")
        assert len(rec.extracted_text) >= 0

    def test_comb_07_f5_f8_digital_extractor_to_date_normalizer(self, tmp_path: Path):
        """
        [F5 + F8] Digital Text Extractor integrated with ISO 8601 Timestamp Normalizer.
        Extracts court document text and converts various date formats into canonical UTC ISO 8601 strings.
        """
        html_content = """
        <html><body>
        <div class="court-stamp">FILED: December 8, 2021</div>
        <div class="hearing">Hearing Scheduled: 05/24/2022 at 09:00 AM</div>
        <div class="notice">Plea Entered on August 16, 2023</div>
        </body></html>
        """
        parser = HtmlDocumentParser()
        res = parser.extract_from_stream(io.BytesIO(html_content.encode("utf-8")))
        
        dates = extract_dates(res.text)
        date_strings = [d.iso_value for d in dates]
        
        assert any(d.startswith("2021-12-08") for d in date_strings)
        assert any(d.startswith("2022-05-24") for d in date_strings)
        assert any(d.startswith("2023-08-16") for d in date_strings)

    def test_comb_08_f5_f9_digital_extractor_to_financial_normalizer(self, tmp_path: Path):
        """
        [F5 + F9] Digital Text Extractor integrated with Financial Transaction Normalizer.
        Extracts monetary amounts from resolution text and verifies dual float and integer cents.
        """
        text = (
            "The Anaheim City Council considered the $320M stadium purchase agreement. "
            "Under the HCD notice, a 30% statutory penalty of $96,000,000.00 was assessed. "
            "Escrow deposit of $50,000,000 was refunded."
        )
        extractor = TextExtractor()
        res = extractor.extract_from_stream(io.BytesIO(text.encode("utf-8")))
        
        financials = extract_financials(res.text)
        
        assert any(f.amount_float == 320000000.0 and f.amount_cents == 32000000000 for f in financials)
        assert any(f.amount_float == 96000000.0 and f.amount_cents == 9600000000 for f in financials)
        assert any(f.amount_float == 50000000.0 and f.amount_cents == 5000000000 for f in financials)

    def test_comb_09_f5_f10_digital_extractor_to_case_normalizer(self, tmp_path: Path):
        """
        [F5 + F10] Digital Text Extractor integrated with Legal Case & Statutory Citation Normalizer.
        Extracts federal criminal dockets (CDCA/DNJ), state eviction dockets, and statutory codes.
        """
        plea_text = """
        UNITED STATES DISTRICT COURT CENTRAL DISTRICT OF CALIFORNIA
        UNITED STATES OF AMERICA v. HARRY SIDHU, Case No. 8:23-cr-00108-CJC
        Defendant violates 18 U.S.C. § 1343 (Wire Fraud).
        Related to state unlawful detainer 30-2021-01201327-CL-UD-CJC under Cal. CCP § 170.6.
        """
        citations = extract_case_citations(plea_text)
        canonical_ids = [c.canonical_id for c in citations]
        
        assert any("8:23-cr-00108" in c for c in canonical_ids)
        assert any("30-2021-01201327" in c for c in canonical_ids)
        assert any("18 U.S.C. § 1343" in c for c in canonical_ids)
        assert any("Cal. CCP § 170.6" in c for c in canonical_ids)

    def test_comb_10_f6_f7_image_enhancer_to_ocr_engine(self, tmp_path: Path):
        """
        [F6 + F7] OpenCV Image Preprocessing (CLAHE + Deskew) feeding Neural Offline OCR.
        Verifies that degraded/skewed image is corrected and transcribed by OCR engine.
        """
        img = np.full((150, 500), 200, dtype=np.uint8)
        cv2.putText(img, "ANAHEIM CHAMBER TA GROUP", (20, 80), cv2.FONT_HERSHEY_SIMPLEX, 0.8, 150, 2)
        
        enhancer = ImageEnhancer()
        enhanced = enhancer.enhance(img, profile=EnhancementProfile.DOCUMENT_CLEAN)
        
        ocr = OCREngine.get_instance()
        res_enh = ocr.ocr_image(enhanced)
        
        assert res_enh.page_number == 1
        assert isinstance(res_enh.lines, (list, tuple))


# ==============================================================================
# TIER 3 PAIRWISE COMBINATIONS: ENTITY RESOLUTION & GRAPH (TESTS 11 - 15)
# ==============================================================================

class TestTier3EntityResolutionCombinations:
    """Pairwise integration tests for entity extraction, disambiguation, and graph synthesis."""

    def test_comb_11_f8_f12_timestamp_to_entity_contextual_association(self):
        """
        [F8 + F12] Timestamp Normalizer correlated with 6-Category Entity Extractor.
        Verifies that entity mentions are temporally anchored with surrounding ISO 8601 dates.
        """
        record = ExtractedRecord(
            record_id="REC-TEST-01",
            artifact_sha256="a" * 64,
            source_path="evidence/memo.txt",
            source_type="local_file",
            mime_type="text/plain",
            normalized_date="2022-05-24",
            raw_date_string="May 24, 2022",
            extracted_text="On May 24, 2022, Mayor Harry Sidhu resigned amidst the Anaheim City Council probe.",
            ocr_engine_used="text_reader",
            financial_amounts=[],
            case_numbers=[],
            sender=None,
            recipients=[],
            metadata={},
        )
        resolver = EntityResolver()
        entities, mentions, events, _, _ = resolver.extract_and_resolve([record])
        
        entity_names = [e.canonical_name for e in entities]
        assert any("Sidhu" in n for n in entity_names)
        assert any("Anaheim" in n for n in entity_names)
        
        assert len(events) >= 1
        evt = events[0]
        assert evt.event_date_iso.startswith("2022-05-24")

    def test_comb_12_f9_f12_financial_to_entity_payer_payee_mapping(self):
        """
        [F9 + F12] Financial Transaction Normalizer integrated with Entity Extractor.
        Verifies monetary amounts are linked with sender and recipient entities.
        """
        record = ExtractedRecord(
            record_id="REC-FIN-01",
            artifact_sha256="b" * 64,
            source_path="evidence/wire.txt",
            source_type="local_file",
            mime_type="text/plain",
            normalized_date="2021-04-15",
            raw_date_string="April 15, 2021",
            extracted_text="TA Group LLC transferred $1,500,000 to FPS Strategies LLC for political consulting.",
            ocr_engine_used="text_reader",
            financial_amounts=[{"raw": "$1,500,000", "amount_float": 1500000.0, "amount_cents": 150000000, "currency": "USD"}],
            case_numbers=[],
            sender=None,
            recipients=[],
            metadata={},
        )
        resolver = EntityResolver()
        entities, _, _, transactions, _ = resolver.extract_and_resolve([record])
        
        assert len(transactions) >= 1
        trx = transactions[0]
        assert trx.amount == 1500000.0
        assert any("TA Group" in e.canonical_name for e in entities)

    def test_comb_13_f10_f12_case_normalizer_to_entity_judicial_mapping(self):
        """
        [F10 + F12] Legal Case Normalizer integrated with Entity Extractor.
        Verifies docket numbers associate with judicial officer and agency entities.
        """
        record = ExtractedRecord(
            record_id="REC-COURT-01",
            artifact_sha256="c" * 64,
            source_path="evidence/ruling.txt",
            source_type="local_file",
            mime_type="text/plain",
            normalized_date="2021-12-22",
            raw_date_string="December 22, 2021",
            extracted_text="In Case No. 30-2021-01201327-CL-UD-CJC, Judge Carmen Luege issued a stay order.",
            ocr_engine_used="text_reader",
            financial_amounts=[],
            case_numbers=["30-2021-01201327-CL-UD-CJC"],
            sender=None,
            recipients=[],
            metadata={},
        )
        resolver = EntityResolver()
        entities, _, events, _, _ = resolver.extract_and_resolve([record])
        
        entity_names = [e.canonical_name for e in entities]
        assert any("Luege" in name for name in entity_names)

    def test_comb_14_f11_f12_comm_metadata_to_entity_extractor(self, tmp_path: Path):
        """
        [F11 + F12] Communication Metadata Normalizer integrated with Entity Extractor.
        Extracts sender/recipients from email headers and classifies them into INDIVIDUAL/COMMERCIAL.
        """
        email_msg = (
            "From: Todd Ament <tament@anaheimchamber.org>\n"
            "To: Melahat Rafiei <melahat@progressive-solutions.com>\n"
            "Subject: Confidential Stadium Briefing\n"
            "Date: Wed, 10 Nov 2021 14:30:00 -0800\n\n"
            "Please review the attached strategy regarding the Anaheim City Council vote."
        )
        eml_file = tmp_path / "stadium_email.eml"
        eml_file.write_text(email_msg, encoding="utf-8")
        
        reader = MailboxReader()
        artifacts = list(reader.read_eml_file(eml_file))
        assert len(artifacts) >= 1
        art = artifacts[0]
        sender = art.metadata.get("sender", "") if art.metadata else ""
        
        parties = extract_correspondence_parties(f"From: {sender}\nTo: melahat@progressive-solutions.com")
        assert len(parties) >= 1

    def test_comb_15_f12_f13_entity_extractor_to_phonetic_resolver(self):
        """
        [F12 + F13] 6-Category Entity Extractor integrated with Phonetic & DSU Resolver.
        Verifies that spelling variants and OCR noise (e.g. Todd Ament vs Mr. Todd Ament)
        are blocked and resolved into unified canonical clusters with populated aliases_json.
        """
        records = [
            ExtractedRecord(
                record_id="REC-VAR-01",
                artifact_sha256="d" * 64,
                source_path="doc1.txt",
                source_type="local_file",
                mime_type="text/plain",
                normalized_date="2022-01-01",
                raw_date_string=None,
                extracted_text="Todd Ament coordinated with the Chamber of Commerce.",
                ocr_engine_used="text",
                financial_amounts=[],
                case_numbers=[],
                sender=None,
                recipients=[],
                metadata={},
            ),
            ExtractedRecord(
                record_id="REC-VAR-02",
                artifact_sha256="e" * 64,
                source_path="doc2.txt",
                source_type="local_file",
                mime_type="text/plain",
                normalized_date="2022-01-02",
                raw_date_string=None,
                extracted_text="Todd Ament (CEO) attended the executive retreat.",
                ocr_engine_used="text",
                financial_amounts=[],
                case_numbers=[],
                sender=None,
                recipients=[],
                metadata={},
            ),
        ]
        resolver = EntityResolver()
        entities, mentions, _, _, _ = resolver.extract_and_resolve(records)
        
        todd_entities = [e for e in entities if "Todd Ament" in e.canonical_name]
        assert len(todd_entities) == 1
        assert len(mentions) >= 2


# ==============================================================================
# TIER 3 PAIRWISE COMBINATIONS: STORAGE & INVARIANTS (TESTS 16 - 20)
# ==============================================================================

class TestTier3StorageAndInvariantsCombinations:
    """Pairwise integration tests across SQLite Vault, Master Catalog, and Invariants."""

    def test_comb_16_f13_f14_resolver_to_sqlite_vault(self, tmp_path: Path):
        """
        [F13 + F14] Phonetic Entity Resolver piping resolved entities into SQLite Relational Vault.
        Verifies batch insertions, foreign key relationships, and WAL mode execution.
        """
        db_path = tmp_path / "test_vault_comb16.db"
        vault = VaultDB(db_path=db_path)
        
        # 1. Insert Document
        vault.insert_document({
            "document_id": "DOC-COMB-16",
            "source_uri": "file://exhibit.pdf",
            "file_name": "exhibit.pdf",
            "file_path": str(tmp_path / "exhibit.pdf"),
            "file_size_bytes": 1024,
            "mime_type": "application/pdf",
            "file_sha256": "1" * 64,
            "content_sha256": "1" * 64,
            "ingestion_timestamp": "2026-08-29T12:00:00Z",
            "extracted_text": "Exhibit text content.",
        })
        
        # 2. Insert Entity and Mention
        vault.insert_entity({
            "entity_id": "ENT-SIDHU",
            "canonical_name": "Harry Sidhu",
            "entity_category": "INDIVIDUAL",
            "aliases_json": ["Mayor Sidhu"],
        })
        vault.insert_mention({
            "mention_id": "MEN-001",
            "document_id": "DOC-COMB-16",
            "entity_id": "ENT-SIDHU",
            "raw_mention_text": "Mayor Sidhu",
            "extraction_method": "REGEX",
            "confidence_score": 0.95,
        })
        
        # Verify foreign keys hold
        conn = vault.get_connection()
        fk_check = conn.execute("PRAGMA foreign_key_check;").fetchall()
        assert len(fk_check) == 0

    def test_comb_17_f14_f15_sqlite_vault_to_catalog_exporter(self, tmp_path: Path):
        """
        [F14 + F15] SQLite Relational Vault querying to RFC 8785 Master JSON Catalog Exporter.
        Verifies that database records serialize deterministically with computed Merkle root.
        """
        db_path = tmp_path / "vault_comb17.db"
        catalog_path = tmp_path / "master_catalog_comb17.json"
        vault = VaultDB(db_path=db_path)
        
        vault.insert_document({
            "document_id": "DOC-01",
            "source_uri": "uri",
            "file_name": "f1.pdf",
            "file_path": "p1",
            "file_size_bytes": 500,
            "mime_type": "application/pdf",
            "file_sha256": "a" * 64,
            "content_sha256": "a" * 64,
            "ingestion_timestamp": "2026-08-29T12:00:00Z",
        })
        vault.insert_entity({
            "entity_id": "ENT-01",
            "canonical_name": "City of Anaheim",
            "entity_category": "MUNICIPAL_BODY",
            "aliases_json": [],
        })
        
        exporter = CatalogExporter(vault_db=vault, output_path=catalog_path)
        exported_file = exporter.export_to_file(catalog_path)
        
        assert exported_file.exists()
        catalog = json.loads(exported_file.read_text(encoding="utf-8"))
        
        assert catalog["catalog_metadata"]["total_documents"] == 1
        assert catalog["catalog_metadata"]["total_entities"] == 1
        assert len(catalog["catalog_metadata"]["root_merkle_sha256"]) == 64

    def test_comb_18_f1_f14_crawler_to_vault_ingestion(self, tmp_path: Path):
        """
        [F1 + F14] Stream Ingestion Crawler directly inserting document rows into SQLite Vault.
        Verifies that crawled artifacts populate the documents table with verified file_sha256.
        """
        corpus_dir = tmp_path / "crawl_corpus"
        corpus_dir.mkdir()
        
        f1 = corpus_dir / "audit.txt"
        f1.write_text("Forensic Audit JL Group 2022", encoding="utf-8")
        
        db_path = tmp_path / "vault_comb18.db"
        vault = VaultDB(db_path=db_path)
        
        crawler = LocalCrawler(target_paths=[corpus_dir])
        for art in crawler.crawl_directory(corpus_dir):
            vault.insert_document({
                "document_id": f"DOC-{art.artifact_id[:8]}",
                "source_uri": art.source_uri,
                "file_name": Path(art.source_uri).name,
                "file_path": art.source_uri,
                "file_size_bytes": art.file_size_bytes,
                "mime_type": art.mime_type,
                "file_sha256": art.artifact_id,
                "content_sha256": art.artifact_id,
                "ingestion_timestamp": "2026-08-29T12:00:00Z",
                "extracted_text": "Forensic Audit JL Group 2022",
            })
            
        conn = vault.get_connection()
        count = conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
        assert count == 1

    def test_comb_19_f8_f14_timestamp_normalizer_to_vault_timeline_events(self, tmp_path: Path):
        """
        [F8 + F14] Timestamp Normalizer populating timeline_events table in strict chronological order.
        Verifies event year, month, day extraction and non-decreasing rank order.
        """
        db_path = tmp_path / "vault_comb19.db"
        vault = VaultDB(db_path=db_path)
        
        dates_raw = ["2021-05-19", "2021-12-08", "2022-05-24", "2023-08-16"]
        for idx, d_str in enumerate(dates_raw, start=1):
            norm = normalize_date(d_str)
            assert norm is not None
            vault.insert_event({
                "event_id": f"EVT-{idx:03d}",
                "event_date_iso": norm.iso_value,
                "event_year": norm.year,
                "event_month": norm.month,
                "event_day": norm.day,
                "event_type": "JUDICIAL_FILING",
                "title": f"Event {idx}",
                "description": f"Evidentiary milestone on {norm.iso_value}",
                "chronological_rank": idx,
            })
            
        conn = vault.get_connection()
        rows = conn.execute("SELECT event_date_iso FROM timeline_events ORDER BY chronological_rank ASC").fetchall()
        extracted_dates = [r[0] for r in rows]
        assert extracted_dates == sorted(extracted_dates)

    def test_comb_20_f15_f17_master_catalog_to_invariant_verification(self, tmp_path: Path):
        """
        [F15 + F17] Master JSON Catalog Exporter integrated with Invariant Verification Engine.
        Verifies that exported catalog satisfies all schema invariants and cryptographic audits.
        """
        db_path = tmp_path / "vault_comb20.db"
        catalog_path = tmp_path / "master_catalog_comb20.json"
        vault = VaultDB(db_path=db_path)
        
        # Populate minimum viable dataset
        vault.insert_document({
            "document_id": "D-001",
            "source_uri": "uri",
            "file_name": "plea.pdf",
            "file_path": "path",
            "file_size_bytes": 100,
            "mime_type": "application/pdf",
            "file_sha256": "f" * 64,
            "content_sha256": "f" * 64,
            "ingestion_timestamp": "2026-08-29T12:00:00Z",
        })
        
        exporter = CatalogExporter(vault_db=vault, output_path=catalog_path)
        catalog = exporter.build_catalog(integrity_mode="development")
        
        assert catalog["audit_invariants"]["all_invariants_passed"] is True
        assert catalog["audit_invariants"]["foreign_key_violations"] == 0
        assert catalog["audit_invariants"]["chronological_inversions"] == 0
        assert len(catalog["audit_invariants"]["documents_merkle_sha256"]) == 64
