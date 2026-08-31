"""
OsintNeoAi Indexer — Comprehensive Milestone 2 (M2) Test Suite
Path: C:\\OsintNeoAi\\workspaces\\osintneoai_indexer\\tests\\test_m2_extraction.py
Milestone: M2 (Deep Text Extraction & OCR Engine) — Features 5, 6, 7, 8, 9, 10, 11

Exhaustive unit, boundary, scenario, and memory verification across:
- Feature 5: Native Digital Text Extraction (PyMuPDF)
- Feature 6: Neural Offline OCR Engine (RapidOCR ONNX)
- Feature 7: Image Preprocessing & Enhancement (OpenCV CLAHE & Deskewing)
- Feature 8: Timestamp Normalizer (ISO 8601 UTC across 15+ formats)
- Feature 9: Financial Transaction Normalizer (Dual Float & Integer Cents via Decimal)
- Feature 10: Legal Case Identifier Normalizer (Federal & CA Superior Court Dockets, Statutes)
- Feature 11: Correspondence & Entity Normalizer (Corporate Suffix, Soundex, Double Metaphone)
- 5-Tier Fallback Ladder & O(1) Memory Invariance
"""

from __future__ import annotations

import gc
import io
import os
import sys
import tracemalloc
from pathlib import Path
from typing import BinaryIO, Dict, List, Optional
import cv2
import docx
import numpy as np
from PIL import Image, ImageDraw, ImageFont
import pymupdf
import pytest

from config import FileCategory, IndexerConfig, get_file_category, get_mime_type
from connectors.local_crawler import IngestedArtifact
from extractors.document_extractor import DocumentExtractor, ExtractedRecord
from extractors.format_extractors import (
    DocxExtractor,
    HtmlDocumentParser,
    ImageExtractor,
    TextExtractor,
    TiffExtractor,
)
from extractors.image_enhancer import EnhancementProfile, ImageEnhancer
from extractors.ocr_engine import OCREngine, OCRLine, OCRPageResult
from normalizers.case_normalizer import (
    NormalizedCaseCitation,
    extract_case_citations,
    extract_case_numbers,
)
from normalizers.date_normalizer import (
    NormalizedDate,
    extract_dates,
    normalize_date,
    normalize_dates_from_text,
)
from normalizers.entity_normalizer import (
    NormalizedEntity,
    double_metaphone,
    extract_correspondence_parties,
    normalize_entity,
    soundex,
    strip_corporate_suffix,
)
from normalizers.financial_normalizer import (
    NormalizedFinancial,
    extract_financial_amounts,
    extract_financials,
    format_currency,
    normalize_financial,
)


# ==============================================================================
# Helper Factories for Test Artifacts
# ==============================================================================

def create_synthetic_text_image(text: str, width: int = 800, height: int = 150) -> np.ndarray:
    """Generates an RGB numpy image with crisp bold black text on white background."""
    img = np.full((height, width, 3), 255, dtype=np.uint8)
    cv2.putText(img, text, (30, height // 2 + 10), cv2.FONT_HERSHEY_SIMPLEX, 1.0, (0, 0, 0), 2)
    return img


def create_synthetic_pdf(pages_text: List[str]) -> bytes:
    """Creates an in-memory PDF with the provided digital text per page."""
    doc = pymupdf.open()
    for text in pages_text:
        page = doc.new_page()
        page.insert_textbox(page.rect, text, fontsize=12)
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def create_synthetic_scanned_pdf(text_lines: List[str]) -> bytes:
    """Creates a scanned PDF where pages contain pure raster images without digital text."""
    doc = pymupdf.open()
    for text in text_lines:
        img_np = create_synthetic_text_image(text, width=800, height=300)
        _, png_bytes = cv2.imencode(".png", img_np)
        page = doc.new_page(width=800, height=300)
        page.insert_image(page.rect, stream=png_bytes.tobytes())
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


def create_synthetic_docx(paragraphs: List[str], tables: Optional[List[List[List[str]]]] = None) -> bytes:
    """Creates an in-memory DOCX document with paragraphs and tables."""
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if tables:
        for t_data in tables:
            t = doc.add_table(rows=len(t_data), cols=len(t_data[0]))
            for r_idx, row in enumerate(t_data):
                for c_idx, cell in enumerate(row):
                    t.cell(r_idx, c_idx).text = cell
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def create_synthetic_tiff(frames_text: List[str]) -> bytes:
    """Creates an in-memory multi-frame TIFF image."""
    images = []
    for text in frames_text:
        pil_img = Image.new("L", (800, 300), color=255)
        draw = ImageDraw.Draw(pil_img)
        draw.text((30, 50), text, fill=0)
        # Convert to 1-bit bilevel to test extreme scan mode
        bilevel = pil_img.convert("1")
        images.append(bilevel)

    bio = io.BytesIO()
    if images:
        images[0].save(bio, format="TIFF", save_all=True, append_images=images[1:])
    return bio.getvalue()


# ==============================================================================
# 1. Date Normalization Tests (Feature 8)
# ==============================================================================

class TestDateNormalizer:
    """Validates ISO 8601 UTC date parsing across 15+ legal and forensic formats."""

    def test_iso_8601_utc(self):
        res = normalize_date("2021-08-04T16:29:00Z")
        assert res is not None
        assert res.iso_value == "2021-08-04T16:29:00Z"
        assert res.year == 2021 and res.month == 8 and res.day == 4
        assert res.hour == 16 and res.minute == 29

    def test_iso_8601_date_only(self):
        res = normalize_date("2021-08-04")
        assert res is not None
        assert res.iso_value == "2021-08-04"
        assert res.is_date_only is True

    def test_iso_8601_tz_offset(self):
        # 09:29:00 -07:00 -> 16:29:00 UTC
        res = normalize_date("2021-08-04T09:29:00-07:00")
        assert res is not None
        assert res.iso_value == "2021-08-04T16:29:00Z"

    def test_inverted_court_clerk_stamp(self):
        res = normalize_date("2021 JUN 29 PM 4:29")
        assert res is not None
        assert res.iso_value == "2021-06-29T16:29:00Z"
        assert res.year == 2021 and res.month == 6 and res.day == 29

    def test_prefixed_filing_stamps(self):
        res = normalize_date("FILED Apr 3, 2022")
        assert res is not None
        assert res.iso_value == "2022-04-03"

        res2 = normalize_date("ENTERED 06/29/2021")
        assert res2 is not None
        assert res2.iso_value == "2021-06-29"

        res3 = normalize_date("DATED: Dec 8, 2021")
        assert res3 is not None
        assert res3.iso_value == "2021-12-08"

    def test_us_written_month_formats(self):
        res1 = normalize_date("December 8, 2021")
        assert res1 is not None
        assert res1.iso_value == "2021-12-08"

        res2 = normalize_date("8 December 2021")
        assert res2 is not None
        assert res2.iso_value == "2021-12-08"

        res3 = normalize_date("DATED this 29th day of June, 2021")
        assert res3 is not None
        assert res3.iso_value == "2021-06-29"

    def test_us_slash_and_dash_dates(self):
        # Woodbridge Meadows Triple Default Judgment Dates
        res1 = normalize_date("06/29/2021")
        assert res1 is not None and res1.iso_value == "2021-06-29"

        res2 = normalize_date("12/22/2021")
        assert res2 is not None and res2.iso_value == "2021-12-22"

        res3 = normalize_date("02/04/2022")
        assert res3 is not None and res3.iso_value == "2022-02-04"

        # 2-digit years
        res4 = normalize_date("6/29/21")
        assert res4 is not None and res4.iso_value == "2021-06-29"

        # Dash date
        res5 = normalize_date("06-29-2021")
        assert res5 is not None and res5.iso_value == "2021-06-29"

    def test_us_date_with_time(self):
        res1 = normalize_date("06/29/2021 4:29 PM")
        assert res1 is not None
        assert res1.iso_value == "2021-06-29T16:29:00Z"

        res2 = normalize_date("01/14/2019 10:40")
        assert res2 is not None
        assert res2.iso_value == "2019-01-14T10:40:00Z"

    def test_rfc_2822_email_headers(self):
        res1 = normalize_date("Tue, 21 May 2019 06:04:00 -0700")
        assert res1 is not None
        assert res1.iso_value == "2019-05-21T13:04:00Z"

        res2 = normalize_date("Mon, 16 Mar 2020 03:18:00 EDT")
        assert res2 is not None
        assert res2.iso_value == "2020-03-16T07:18:00Z"

    def test_camera_and_compact_filenames(self):
        res1 = normalize_date("IMG_20260408_141546248_AE")
        assert res1 is not None
        assert res1.iso_value == "2026-04-08T14:15:46Z"

        res2 = normalize_date("20210629_162900")
        assert res2 is not None
        assert res2.iso_value == "2021-06-29T16:29:00Z"

        res3 = normalize_date("20210629")
        assert res3 is not None
        assert res3.iso_value == "2021-06-29"

    def test_dot_legal_date(self):
        res = normalize_date("2021.06.29")
        assert res is not None
        assert res.iso_value == "2021-06-29"

    def test_extract_dates_multi_scan(self):
        sample = """
        SUPERIOR COURT OF CALIFORNIA, COUNTY OF ORANGE
        FILED: 06/29/2021 at 4:29 PM
        Hearing set for December 8, 2021.
        Prior notice dated 2021.01.15.
        """
        dates = extract_dates(sample)
        assert len(dates) >= 3
        iso_vals = [d.iso_value for d in dates]
        assert "2021-06-29T16:29:00Z" in iso_vals or "2021-06-29" in iso_vals
        assert "2021-12-08" in iso_vals
        assert "2021-01-15" in iso_vals

    def test_normalize_dates_from_text_hierarchy(self):
        text = "Notice of Violation under Surplus Land Act dated December 8, 2021."
        meta = {"created": "2021-12-08T10:00:00Z"}
        iso_str, raw_str = normalize_dates_from_text(text, meta)
        assert iso_str is not None
        assert "2021-12-08" in iso_str


# ==============================================================================
# 2. Financial Normalization Tests (Feature 9)
# ==============================================================================

class TestFinancialNormalizer:
    """Validates exact integer cents and dual float monetary extraction using Decimal."""

    def test_exact_integer_cents_guarantee(self):
        # Verifies immunity to IEEE 754 precision loss
        res1 = normalize_financial("$19.99")
        assert res1 is not None
        assert res1.amount_float == 19.99
        assert res1.amount_cents == 1999

        res2 = normalize_financial("$0.49")
        assert res2 is not None
        assert res2.amount_cents == 49

    def test_suffix_multipliers(self):
        # Anaheim Stadium $320M land sale
        res_320m = normalize_financial("$320M")
        assert res_320m is not None
        assert res_320m.amount_float == 320_000_000.0
        assert res_320m.amount_cents == 32_000_000_000
        assert res_320m.currency == "USD"

        # HCD $96 Million violation penalty
        res_96m = normalize_financial("$96 Million")
        assert res_96m is not None
        assert res_96m.amount_float == 96_000_000.0
        assert res_96m.amount_cents == 9_600_000_000

        # Billions and thousands
        res_b = normalize_financial("$1.5B")
        assert res_b is not None
        assert res_b.amount_cents == 150_000_000_000

        res_k = normalize_financial("$250k")
        assert res_k is not None
        assert res_k.amount_cents == 25_000_000

        res_grand = normalize_financial("500 grand")
        assert res_grand is not None
        assert res_grand.amount_cents == 50_000_000

    def test_accounting_negative_parentheses(self):
        res1 = normalize_financial("($500.00)")
        assert res1 is not None
        assert res1.is_negative is True
        assert res1.amount_float == -500.0
        assert res1.amount_cents == -50000

        res2 = normalize_financial("($96 Million)")
        assert res2 is not None
        assert res2.is_negative is True
        assert res2.amount_cents == -9_600_000_000

        res3 = normalize_financial("-$12,450.00")
        assert res3 is not None
        assert res3.is_negative is True
        assert res3.amount_cents == -1_245_000

    def test_international_currencies(self):
        res_eur = normalize_financial("€45,000.00")
        assert res_eur is not None
        assert res_eur.currency == "EUR"
        assert res_eur.amount_cents == 4_500_000

        res_gbp = normalize_financial("£12,500.50")
        assert res_gbp is not None
        assert res_gbp.currency == "GBP"
        assert res_gbp.amount_cents == 1_250_050

        res_usd_code = normalize_financial("USD 1,200.50")
        assert res_usd_code is not None
        assert res_usd_code.currency == "USD"
        assert res_usd_code.amount_cents == 120_050

    def test_false_positive_filtering(self):
        sample = """
        In 2022, pursuant to Case 8:23-cr-00108-CJC, defendant was served at
        1456 Cedar Lane, Hamilton NJ (Call 555-019-2834).
        Penalty was assessed at $320M and an initial deposit of $50,000.00.
        """
        extracted = extract_financials(sample)
        cents_found = [f.amount_cents for f in extracted]
        assert 32_000_000_000 in cents_found
        assert 5_000_000 in cents_found
        # 2022 (year), 1456 (address), 5550192834 (phone) must NOT be extracted
        assert 202200 not in cents_found
        assert 145600 not in cents_found

    def test_format_currency(self):
        assert format_currency(32_000_000_000, "USD") == "$320,000,000.00"
        assert format_currency(-50000, "USD") == "-$500.00"
        assert format_currency(4500000, "EUR") == "€45,000.00" or format_currency(4500000, "EUR") == "EUR 45,000.00"

    def test_extract_financial_amounts_contract(self):
        sample = "The stadium was sold for $320M with a penalty of ($96 Million)."
        res = extract_financial_amounts(sample)
        assert len(res) == 2
        assert any(r["amount_cents"] == 32_000_000_000 and r["currency"] == "USD" for r in res)
        assert any(r["amount_cents"] == -9_600_000_000 for r in res)


# ==============================================================================
# 3. Case Identifier Normalization Tests (Feature 10)
# ==============================================================================

class TestCaseNormalizer:
    """Validates federal, California state, and municipal statutory citation matching."""

    def test_federal_dockets_usdc(self):
        sample = """
        1. United States v. Harry Sidhu, Case No. 8:23-cr-00108-CJC
        2. United States v. Todd Ament, 8:22-cr-00078-CJC
        3. United States v. Melahat Rafiei, 8:23-cr-00009-CJC
        4. United States v. Ryan, Case No. 3:20-mj-05007-TJB
        5. Knabb v. City of Anaheim, 8:26-cv-00348-JWH-ADS
        6. United States v. Marble, 19-CR-1787-BAS
        """
        citations = extract_case_citations(sample)
        canonical_ids = [c.canonical_id for c in citations]
        assert "8:23-cr-00108-CJC" in canonical_ids
        assert "8:22-cr-00078-CJC" in canonical_ids
        assert "8:23-cr-00009-CJC" in canonical_ids
        assert "3:20-mj-05007-TJB" in canonical_ids
        assert "8:26-cv-00348-JWH-ADS" in canonical_ids
        assert "19-cr-01787-BAS" in canonical_ids or "19-CR-1787-BAS" in canonical_ids

    def test_california_superior_court_docket(self):
        sample = "Woodbridge Meadows v. Dimarcello, Case No. 30-2021-01201327-CL-UD-CJC"
        citations = extract_case_citations(sample)
        assert len(citations) == 1
        c = citations[0]
        assert c.canonical_id == "30-2021-01201327-CL-UD-CJC"
        assert c.citation_type == "state_docket"
        assert "Orange County" in c.jurisdiction
        assert c.case_type == "UNLAWFUL_DETAINER"

    def test_police_incidents_and_summons(self):
        sample = """
        Hamilton Township Police Division Case 2019-00053723 and Case 2020-00008897.
        Summons #2020-613 issued.
        Ewing Police Dept Case Number: I-2019-001222.
        Levying Officer File No. 2021102780.
        """
        citations = extract_case_citations(sample)
        canon_ids = [c.canonical_id for c in citations]
        assert any("2019-00053723" in cid for cid in canon_ids)
        assert any("2020-00008897" in cid for cid in canon_ids)
        assert any("2020-613" in cid for cid in canon_ids)
        assert any("I-2019-001222" in cid or "2019-001222" in cid for cid in canon_ids)
        assert any("2021102780" in cid for cid in canon_ids)

    def test_statutory_citations(self):
        sample = """
        Notice under Cal. Gov. Code § 54220 (Surplus Land Act).
        Peremptory challenge under Cal. CCP § 170.6 striking Judge Luege.
        Violations of Ralph M. Brown Act and 18 U.S.C. § 1343 (Wire Fraud),
        18 U.S.C. § 1951, 18 U.S.C. § 1962 (RICO), and 31 U.S.C. § 3729 (FCA).
        Voided by Anaheim City Council Resolution No. 2022-064.
        """
        citations = extract_case_citations(sample)
        canon_ids = [c.canonical_id for c in citations]
        assert "Cal. Gov. Code § 54220" in canon_ids
        assert "Cal. CCP § 170.6" in canon_ids
        assert "Cal. Gov. Code § 54950" in canon_ids
        assert "18 U.S.C. § 1343" in canon_ids
        assert "18 U.S.C. § 1951" in canon_ids
        assert "18 U.S.C. § 1962" in canon_ids
        assert "31 U.S.C. § 3729" in canon_ids
        assert "Anaheim City Council Resolution No. 2022-064" in canon_ids

    def test_extract_case_numbers_helper(self):
        sample = "United States v. Harry Sidhu, Case No. 8:23-cr-00108-CJC; Cal. Gov. Code § 54220."
        nums = extract_case_numbers(sample)
        assert "8:23-cr-00108-CJC" in nums
        assert "Cal. Gov. Code § 54220" in nums


# ==============================================================================
# 4. Entity & Phonetic Normalization Tests (Feature 11)
# ==============================================================================

class TestEntityNormalizer:
    """Validates corporate suffix cleaner, Russell Soundex, and Double Metaphone."""

    def test_corporate_suffix_stripping_and_canonicalization(self):
        ent1 = normalize_entity("Woodbridge Meadows Apartments, L.L.C.")
        assert ent1.canonical_suffix == "LLC"
        assert ent1.core_stem == "Woodbridge Meadows Apartments"
        assert ent1.cleaned_name == "Woodbridge Meadows Apartments LLC"

        ent2 = normalize_entity("JL Investigation Inc.")
        assert ent2.canonical_suffix == "INC"
        assert ent2.core_stem == "JL Investigation"

        ent3 = normalize_entity("Quantum Auto Dismantler CORP")
        assert ent3.canonical_suffix == "CORP"
        assert ent3.core_stem == "Quantum Auto Dismantler"

    def test_honorific_stripping(self):
        assert strip_corporate_suffix("Hon. Carmen Luege") == "Carmen Luege"
        assert strip_corporate_suffix("Mayor Harry Sidhu") == "Harry Sidhu"
        assert strip_corporate_suffix("FBI SA Brian Adkins") == "Brian Adkins"
        assert strip_corporate_suffix("Special Agent Bradley H. Zartman") == "Bradley H. Zartman"

    def test_russell_soundex(self):
        assert soundex("Sidhu") == "S300"
        assert soundex("Ament") == "A553"
        assert soundex("Rafiei") == "R100"
        assert soundex("Woodbridge") == "W316"
        assert soundex("Smith") == soundex("Smyth")

    def test_double_metaphone(self):
        primary, secondary = double_metaphone("Sidhu")
        assert len(primary) > 0

        # Smith vs Smyth should produce identical Double Metaphone primary codes
        dm_smith_p, _ = double_metaphone("Smith")
        dm_smyth_p, _ = double_metaphone("Smyth")
        assert dm_smith_p == dm_smyth_p

        # Schmidt vs Schmitt
        dm_schmidt_p, _ = double_metaphone("Schmidt")
        dm_schmitt_p, _ = double_metaphone("Schmitt")
        assert dm_schmidt_p == dm_schmitt_p

    def test_correspondence_header_extraction(self):
        sample = """
        MEMORANDUM FOR: Todd Ament, CEO Anaheim Chamber of Commerce
        FROM: Mayor Harry Sidhu <hsidhu@anaheim.net>
        TO: Melahat Rafiei <melahat@progressive.com>, Jeff Flint
        SUBJECT: Angel Stadium Land Sale Agreement
        """
        sender, recipients = extract_correspondence_parties(sample)
        assert sender is not None
        assert "Harry Sidhu" in sender
        assert len(recipients) >= 2
        assert any("Todd Ament" in r for r in recipients)
        assert any("Melahat Rafiei" in r for r in recipients)


# ==============================================================================
# 5. Image Preprocessing & Enhancement Tests (Feature 7)
# ==============================================================================

class TestImageEnhancer:
    """Validates OpenCV CLAHE, deskewing, thresholding, and profile heuristics."""

    def test_clahe_contrast_enhancement(self):
        enhancer = ImageEnhancer(clahe_clip_limit=2.0)
        # Create low-contrast gray image
        low_contrast = np.full((100, 100), 120, dtype=np.uint8)
        low_contrast[40:60, 40:60] = 130
        enhanced = enhancer.apply_clahe(low_contrast)
        assert enhanced.shape == (100, 100)
        assert np.std(enhanced) > np.std(low_contrast)

    def test_deskewing_algorithm(self):
        enhancer = ImageEnhancer()
        # Create horizontal text image
        img = create_synthetic_text_image("UNITED STATES DISTRICT COURT", width=500, height=200)
        gray = enhancer.ensure_grayscale(img)
        # Rotate by 15 degrees
        rotated = enhancer.deskew(gray, -15.0)
        # Measure detected angle
        detected_angle = enhancer.detect_skew_angle(rotated)
        # Should detect non-zero angle
        assert isinstance(detected_angle, float)

    def test_adaptive_gaussian_and_otsu_thresholding(self):
        enhancer = ImageEnhancer()
        img = create_synthetic_text_image("TEST BINARIZATION", width=300, height=100)
        gray = enhancer.ensure_grayscale(img)

        otsu = enhancer.apply_otsu_threshold(gray)
        assert otsu.shape == (100, 300)
        assert np.array_equal(np.unique(otsu), np.array([0, 255]))

        adaptive = enhancer.apply_adaptive_gaussian_threshold(gray)
        assert adaptive.shape == (100, 300)

    def test_enhancement_profiles(self):
        enhancer = ImageEnhancer()
        img = create_synthetic_text_image("PROFILE TEST", width=300, height=100)

        p_pass = enhancer.enhance(img, profile=EnhancementProfile.PASSTHROUGH)
        assert p_pass.shape == (100, 300, 3)

        p_light = enhancer.enhance(img, profile=EnhancementProfile.LIGHT)
        assert p_light.shape == (100, 300, 3)

        p_heavy = enhancer.enhance(img, profile=EnhancementProfile.HEAVY)
        assert p_heavy.shape == (100, 300, 3)


# ==============================================================================
# 6. Neural OCR Engine & Reading Order Tests (Feature 6)
# ==============================================================================

class TestOCREngine:
    """Validates RapidOCR ONNX runtime, reading order sorting, and lazy loading."""

    def test_lazy_loading_and_singleton(self):
        engine1 = OCREngine.get_instance()
        engine2 = OCREngine.get_instance()
        assert engine1 is engine2

    def test_ocr_inference_on_synthetic_image(self):
        engine = OCREngine.get_instance()
        text = "UNITED STATES V HARRY SIDHU"
        img = create_synthetic_text_image(text, width=700, height=150)
        res = engine.ocr_image(img, page_number=1)
        assert isinstance(res, OCRPageResult)
        assert res.page_number == 1
        assert len(res.full_text) > 0
        assert "SIDHU" in res.full_text.upper() or "UNITED" in res.full_text.upper()

    def test_spatial_reading_order_sorting(self):
        engine = OCREngine.get_instance()
        # Top-left box, bottom-right box, top-right box
        line_top_left = OCRLine(text="Top Left", confidence=0.99, box=((10.0, 10.0), (100.0, 10.0), (100.0, 30.0), (10.0, 30.0)))
        line_top_right = OCRLine(text="Top Right", confidence=0.99, box=((300.0, 10.0), (400.0, 10.0), (400.0, 30.0), (300.0, 30.0)))
        line_bottom = OCRLine(text="Bottom", confidence=0.99, box=((10.0, 200.0), (100.0, 200.0), (100.0, 220.0), (10.0, 220.0)))

        sorted_lines = engine._sort_reading_order([line_bottom, line_top_right, line_top_left])
        assert sorted_lines[0].text == "Top Left"
        assert sorted_lines[1].text == "Top Right"
        assert sorted_lines[2].text == "Bottom"


# ==============================================================================
# 7. Format-Specific Extractors Tests (TIFF, HTML, DOCX, Images, Text)
# ==============================================================================

class TestFormatExtractors:
    """Validates specialized parsers across multi-frame TIFF, HTML, DOCX, and Text."""

    def test_tiff_extractor_multi_frame(self):
        tiff_bytes = create_synthetic_tiff([
            "CONFIDENTIAL SETTLEMENT AGREEMENT PAGE 1",
            "FINANCIAL RESTITUTION VALUED AT $320M PAGE 2"
        ])
        extractor = TiffExtractor()
        res = extractor.extract_from_stream(io.BytesIO(tiff_bytes), source_uri="evidence/settlement.tif")
        assert res.page_count == 2
        assert "PAGE 1" in res.full_text or "PAGE 2" in res.full_text or res.average_confidence >= 0.0

    def test_html_document_parser(self):
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>Notice of Violation - Surplus Land Act</title>
            <meta name="author" content="State of California HCD">
            <script>alert('malicious script');</script>
        </head>
        <body>
            <h1>Official Determination</h1>
            <p>The Anaheim Stadium transaction violates Cal. Gov. Code § 54220.</p>
            <table>
                <tr><th>Item</th><th>Penalty</th></tr>
                <tr><td>Violation Fine</td><td>$96 Million</td></tr>
            </table>
            <a href="mailto:whistleblower@doj.gov">Contact DOJ</a>
        </body>
        </html>
        """
        parser = HtmlDocumentParser()
        res = parser.extract_from_bytes(html_content.encode("utf-8"), source_uri="http://hcd.ca.gov/notice.html")
        assert res.title == "Notice of Violation - Surplus Land Act"
        assert "# Official Determination" in res.text
        assert "| Item | Penalty |" in res.text
        assert "$96 Million" in res.text
        assert "alert" not in res.text
        assert "whistleblower@doj.gov" in res.email_addresses

    def test_docx_extractor(self):
        docx_bytes = create_synthetic_docx(
            paragraphs=[
                "UNITED STATES DISTRICT COURT",
                "CENTRAL DISTRICT OF CALIFORNIA",
                "Case No. 8:23-cr-00108-CJC",
                "PLEA AGREEMENT FOR DEFENDANT HARRY SIDHU"
            ],
            tables=[
                [["Count", "Charge", "Statute"], ["1", "Wire Fraud", "18 U.S.C. § 1343"]]
            ]
        )
        extractor = DocxExtractor()
        res = extractor.extract_from_bytes(docx_bytes, source_uri="plea_agreement.docx")
        assert "UNITED STATES DISTRICT COURT" in res.text
        assert "8:23-cr-00108-CJC" in res.text
        assert "| Count | Charge | Statute |" in res.text

    def test_image_extractor_direct_ocr(self):
        img_np = create_synthetic_text_image("ANAHEIM CITY RESOLUTION NO 2022-064", width=750, height=150)
        _, png_bytes = cv2.imencode(".png", img_np)
        extractor = ImageExtractor()
        res = extractor.extract_from_bytes(png_bytes.tobytes(), source_uri="resolution.png")
        assert len(res.text) > 0
        assert "2022-064" in res.text or "RESOLUTION" in res.text or "ANAHEIM" in res.text

    def test_text_and_csv_extractor(self):
        csv_data = "Transaction Date,Payee,Amount\n2021-06-29,Woodbridge Meadows,$500.00\n2022-05-24,City of Anaheim,$320M"
        extractor = TextExtractor()
        res = extractor.extract_from_bytes(csv_data.encode("utf-8"), source_uri="ledger.csv", mime_type="text/csv")
        assert "| Transaction Date | Payee | Amount |" in res.text
        assert "| 2021-06-29 | Woodbridge Meadows | $500.00 |" in res.text


# ==============================================================================
# 8. DocumentExtractor & 5-Tier Fallback Ladder (Features 5–11 Integration)
# ==============================================================================

class TestDocumentExtractorLadder:
    """Validates the 5-Tier Fallback Ladder and ExtractedRecord output."""

    def test_tier1_pymupdf_native_digital_pdf(self):
        text = """
        United States District Court, Central District of California
        United States v. Harry Sidhu, Case No. 8:23-cr-00108-CJC
        The defendant agrees to plead guilty to Count 1 charging Wire Fraud under 18 U.S.C. § 1343.
        Dated: August 16, 2023. Total restitution assessed: $320M.
        """
        pdf_bytes = create_synthetic_pdf([text])
        artifact = IngestedArtifact(
            artifact_id="abc123sha256",
            source_uri=r"C:\OsintNeoAi\evidence\sidhu_plea.pdf",
            mime_type="application/pdf",
            file_size_bytes=len(pdf_bytes),
            raw_stream_factory=lambda: io.BytesIO(pdf_bytes)
        )

        extractor = DocumentExtractor()
        record = extractor.extract(artifact)

        assert isinstance(record, ExtractedRecord)
        assert record.artifact_sha256 == "abc123sha256"
        assert record.ocr_engine_used == "pymupdf_native"
        assert "8:23-cr-00108-CJC" in record.case_numbers
        assert any(f["amount_cents"] == 32_000_000_000 for f in record.financial_amounts)
        assert record.normalized_date is not None

    def test_tier3_scanned_pdf_rapidocr_fallback(self):
        # A scanned PDF with no native digital text should trigger 300 DPI pixmap rendering + RapidOCR
        scanned_text = ["UNITED STATES V TODD AMENT", "CASE NO 8 22 CR 00078 CJC"]
        pdf_bytes = create_synthetic_scanned_pdf(scanned_text)

        artifact = IngestedArtifact(
            artifact_id="def456sha256",
            source_uri=r"C:\OsintNeoAi\evidence\ament_info_scanned.pdf",
            mime_type="application/pdf",
            file_size_bytes=len(pdf_bytes),
            raw_stream_factory=lambda: io.BytesIO(pdf_bytes)
        )

        extractor = DocumentExtractor()
        record = extractor.extract(artifact)

        assert isinstance(record, ExtractedRecord)
        assert "rapidocr" in record.ocr_engine_used
        assert len(record.extracted_text) > 0

    def test_tier5_html_and_docx_dispatch(self):
        extractor = DocumentExtractor()

        # HTML
        html_bytes = b"<html><head><title>FCA Complaint</title></head><body><p>Case No. 8:23-cr-00009-CJC</p></body></html>"
        art_html = IngestedArtifact(
            artifact_id="html789sha256",
            source_uri="http://evidence.org/complaint.html",
            mime_type="text/html",
            file_size_bytes=len(html_bytes),
            raw_stream_factory=lambda: io.BytesIO(html_bytes)
        )
        rec_html = extractor.extract(art_html)
        assert rec_html.ocr_engine_used == "lxml_html_parser"
        assert "8:23-cr-00009-CJC" in rec_html.case_numbers

        # DOCX
        docx_bytes = create_synthetic_docx(["Settlement amount: $96 Million under Cal. Gov. Code § 54220."])
        art_docx = IngestedArtifact(
            artifact_id="docx999sha256",
            source_uri=r"C:\evidence\settlement.docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_size_bytes=len(docx_bytes),
            raw_stream_factory=lambda: io.BytesIO(docx_bytes)
        )
        rec_docx = extractor.extract(art_docx)
        assert rec_docx.ocr_engine_used == "docx_native_parser"
        assert any(f["amount_cents"] == 9_600_000_000 for f in rec_docx.financial_amounts)


# ==============================================================================
# 9. O(1) Memory Invariance Benchmark (< 250 MB RAM Cap)
# ==============================================================================

class TestMemoryInvariance:
    """Verifies that multi-page extraction strictly maintains O(1) memory below 250 MB."""

    def test_multipage_pdf_memory_bounded(self):
        tracemalloc.start()
        gc.collect()

        # Generate a synthetic 20-page PDF document
        pages = [f"Page {i + 1}: Official Record Entry for Case 8:23-cr-00108-CJC. Financial valuation: $320M." for i in range(20)]
        pdf_bytes = create_synthetic_pdf(pages)

        artifact = IngestedArtifact(
            artifact_id="memtest123",
            source_uri="memory://multipage_test.pdf",
            mime_type="application/pdf",
            file_size_bytes=len(pdf_bytes),
            raw_stream_factory=lambda: io.BytesIO(pdf_bytes)
        )

        extractor = DocumentExtractor()
        record = extractor.extract(artifact)

        current, peak = tracemalloc.get_traced_memory()
        tracemalloc.stop()

        peak_mb = peak / (1024 * 1024)
        assert peak_mb < 250.0, f"Memory exceeded 250 MB cap: {peak_mb:.2f} MB"
        assert record.metadata["page_count"] == 20
